# appip/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.core.paginator import Paginator
from django.contrib import messages
from django.db.models import Avg, Sum, Count, Q
from datetime import datetime, timedelta
import json
from .models import *
from django.contrib.auth.hashers import make_password, check_password
from django.utils import timezone
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
import io
from docx import Document
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_GET
from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from .serializers import *
import logging

from decimal import Decimal, InvalidOperation

logger = logging.getLogger(__name__)

# ==================== ОСНОВНЫЕ СТРАНИЦЫ ====================

def home(request):
    """Главная страница"""
    # Используем только необходимые поля
    popular_products = Products.objects.filter(is_active=True).only(
        'id_product', 'title', 'price', 'rating', 'main_image_url'
    ).order_by('-rating')[:8]
    
    new_products = Products.objects.filter(is_active=True).only(
        'id_product', 'title', 'price', 'rating', 'main_image_url', 'created_at'
    ).order_by('-created_at')[:8]
    
    categories = Categories.objects.all()[:6]
    
    return render(request, 'home.html', {
        'popular_products': popular_products,
        'new_products': new_products,
        'categories': categories
    })

def products_list(request):
    """Список товаров"""
    from django.db.models import Avg, Count, Q, OuterRef, Subquery
    from django.db.models.functions import Coalesce
    
    products_list = Products.objects.filter(is_active=True).select_related('seller', 'category', 'product_type')
    
    # Подзапрос для получения среднего рейтинга продавца
    seller_avg_rating = SellerReviews.objects.filter(
        seller=OuterRef('seller_id')
    ).values('seller').annotate(
        avg=Coalesce(Avg('rating'), 0.0)
    ).values('avg')
    
    # Подзапрос для получения количества отзывов продавца
    seller_reviews_count = SellerReviews.objects.filter(
        seller=OuterRef('seller_id')
    ).values('seller').annotate(
        count=Count('id_sellerreview')
    ).values('count')
    
    # Аннотируем товары рейтингом продавца и количеством отзывов на продавца
    products_list = products_list.annotate(
        avg_seller_rating=Coalesce(Subquery(seller_avg_rating), 0.0),
        seller_reviews_count=Coalesce(Subquery(seller_reviews_count), 0)
    )
    
    categories = Categories.objects.all()
    
    # Фильтрация
    category_id = request.GET.get('category')
    if category_id:
        products_list = products_list.filter(category_id=category_id)
    
    product_type = request.GET.get('type')
    if product_type:
        products_list = products_list.filter(product_type_id=product_type)
    
    min_price = request.GET.get('min_price')
    max_price = request.GET.get('max_price')
    if min_price:
        products_list = products_list.filter(price__gte=min_price)
    if max_price:
        products_list = products_list.filter(price__lte=max_price)
    
    # Поиск
    search_query = request.GET.get('q')
    if search_query:
        products_list = products_list.filter(
            Q(title__icontains=search_query) | 
            Q(description__icontains=search_query)
        )
    
    # Сортировка с учетом рейтинга продавца
    sort_by = request.GET.get('sort', 'popular')
    if sort_by == 'price_asc':
        products_list = products_list.order_by('price')
    elif sort_by == 'price_desc':
        products_list = products_list.order_by('-price')
    elif sort_by == 'new':
        products_list = products_list.order_by('-created_at')
    elif sort_by == 'rating':
        # Сортировка по рейтингу продавца
        products_list = products_list.order_by('-avg_seller_rating', '-seller_reviews_count')
    else:  # popular
        # "Популярность" по рейтингу продавца
        products_list = products_list.order_by('-avg_seller_rating', '-seller_reviews_count', '-created_at')
    
    # Пагинация
    from django.core.paginator import Paginator
    paginator = Paginator(products_list, 120)
    page_number = request.GET.get('page')
    products = paginator.get_page(page_number)
    
    return render(request, 'products/list.html', {
        'products': products,
        'categories': categories,
        'product_types': ProductTypes.objects.all(),
        'current_sort': sort_by,
        'search_query': search_query or ''
    })





def admin_roles(request):
    """Страница управления ролями"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    user = Users.objects.get(id_user=user_id)


    
    # Только админ может управлять ролями
    if user.role_id != 1:  # Только администратор
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    roles = Roles.objects.all()
    
    return render(request, 'admin/roles.html', {
        'roles': roles,
        'is_admin': user.role_id == 1
    })

@csrf_exempt
@require_POST
def create_role(request):
    """Создание новой роли"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = Users.objects.get(id_user=user_id)
    
    # Только админ может создавать роли
    if user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        data = json.loads(request.body)
        role_name = data.get('role_name')
        
        if not role_name:
            return JsonResponse({'error': 'Название роли обязательно'}, status=400)
        
        if Roles.objects.filter(role_name=role_name).exists():
            return JsonResponse({'error': 'Роль с таким названием уже существует'}, status=400)
        
        role = Roles.objects.create(role_name=role_name)
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='admin_action',
            description=f'Создана новая роль: {role_name}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Роль успешно создана',
            'role': {
                'id': role.id_role,
                'name': role.role_name
            }
        })
        
    except Exception as e:
        logger.error(f'Create role error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def edit_role(request, role_id):
    """Редактирование роли"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = Users.objects.get(id_user=user_id)
    
    # Только админ может редактировать роли
    if user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        role = get_object_or_404(Roles, id_role=role_id)
        data = json.loads(request.body)
        role_name = data.get('role_name')
        
        if not role_name:
            return JsonResponse({'error': 'Название роли обязательно'}, status=400)
        
        # Проверяем уникальность (кроме текущей роли)
        if Roles.objects.filter(role_name=role_name).exclude(id_role=role_id).exists():
            return JsonResponse({'error': 'Роль с таким названием уже существует'}, status=400)
        
        old_name = role.role_name
        role.role_name = role_name
        role.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='admin_action',
            description=f'Изменена роль: {old_name} -> {role_name}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Роль успешно обновлена'
        })
        
    except Exception as e:
        logger.error(f'Edit role error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def delete_role(request, role_id):
    """Удаление роли"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = Users.objects.get(id_user=user_id)
    
    # Только админ может удалять роли
    if user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        role = get_object_or_404(Roles, id_role=role_id)
        
        # Проверяем, используется ли роль
        user_count = Users.objects.filter(role=role).count()
        if user_count > 0:
            return JsonResponse({
                'error': f'Невозможно удалить роль, так как она используется {user_count} пользователями'
            }, status=400)
        
        role_name = role.role_name
        role.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='admin_action',
            description=f'Удалена роль: {role_name}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Роль успешно удалена'
        })
        
    except Exception as e:
        logger.error(f'Delete role error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)





def admin_users(request):
    """Страница управления пользователями"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут управлять пользователями
    if current_user.role_id != 1:  # Только администратор
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    # Фильтрация и поиск
    search_query = request.GET.get('search', '')
    role_filter = request.GET.get('role', '')
    status_filter = request.GET.get('status', '')
    
    users = Users.objects.all().select_related('role').order_by('-registration_date')
    
    # Применяем фильтры
    if search_query:
        users = users.filter(
            Q(login__icontains=search_query) |
            Q(firstname__icontains=search_query) |
            Q(surname__icontains=search_query)
        )
    
    if role_filter:
        users = users.filter(role_id=role_filter)
    
    if status_filter:
        if status_filter == 'active':
            users = users.filter(is_active=True)
        elif status_filter == 'inactive':
            users = users.filter(is_active=False)
    
    # Получаем все роли для фильтра
    roles = Roles.objects.all()
    
    # Пагинация
    paginator = Paginator(users, 20)  # 20 пользователей на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'admin/users.html', {
        'users': page_obj,
        'page_obj': page_obj,
        'roles': roles,
        'search_query': search_query,
        'role_filter': role_filter,
        'status_filter': status_filter,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

@csrf_exempt
@require_POST
def create_user(request):
    """Создание нового пользователя (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ может создавать пользователей
    if current_user.role_id != 1:  # Только администратор
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        login = data.get('login')
        password = data.get('password')
        firstname = data.get('firstname')
        surname = data.get('surname')
        role_id = data.get('role_id')
        
        # Валидация
        if not all([login, password, firstname, surname, role_id]):
            return JsonResponse({'error': 'Все поля обязательны'}, status=400)
        
        if Users.objects.filter(login=login).exists():
            return JsonResponse({'error': 'Пользователь с таким email уже существует'}, status=400)
        
        if len(password) < 6:
            return JsonResponse({'error': 'Пароль должен содержать минимум 6 символов'}, status=400)
        
        
        
        # Создаем пользователя
        user = Users.objects.create(
            login=login,
            firstname=firstname,
            surname=surname,
            role_id=role_id,
            is_active=True
        )
        user.set_password(password)
        user.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Создан новый пользователь: {login} ({firstname} {surname})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Пользователь успешно создан',
            'user': {
                'id': user.id_user,
                'login': user.login,
                'name': f"{user.firstname} {user.surname}",
                'role': user.role.role_name,
                'registration_date': user.registration_date.strftime('%d.%m.%Y %H:%M'),
                'is_active': user.is_active
            }
        })
        
    except Exception as e:
        logger.error(f'Create user error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def edit_user(request, user_id):
    """Редактирование пользователя (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    current_user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=current_user_id)
    
    # Только админ и менеджер могут редактировать пользователей
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        user = get_object_or_404(Users, id_user=user_id)
        
        # Менеджер не может редактировать админов и других менеджеров
        if current_user.role_id == 3 and user.role_id in [1, 3]:
            return JsonResponse({'error': 'Менеджер не может редактировать администраторов и других менеджеров'}, status=403)
        
        data = json.loads(request.body)
        
        firstname = data.get('firstname')
        surname = data.get('surname')
        role_id = data.get('role_id')
        
        # Валидация
        if not all([firstname, surname, role_id]):
            return JsonResponse({'error': 'Все поля обязательны'}, status=400)
        
        # Менеджер не может менять роль на админа
        if current_user.role_id == 3 and int(role_id) == 1:
            return JsonResponse({'error': 'Менеджер не может назначать роль администратора'}, status=403)
        
        # Менеджер не может менять роль админам и менеджерам
        if current_user.role_id == 3 and user.role_id in [1, 3] and int(role_id) != user.role_id:
            return JsonResponse({'error': 'Менеджер не может изменять роли администраторов и менеджеров'}, status=403)
        
        old_name = f"{user.firstname} {user.surname}"
        old_role = user.role.role_name
        
        user.firstname = firstname
        user.surname = surname
        user.role_id = role_id
        user.save()
        
        new_role = Roles.objects.get(id_role=role_id).role_name
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Отредактирован пользователь {user.login}: {old_name} ({old_role}) -> {firstname} {surname} ({new_role})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Пользователь успешно обновлен'
        })
        
    except Exception as e:
        logger.error(f'Edit user error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def delete_user(request, user_id):
    """Удаление пользователя (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    current_user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=current_user_id)
    
    # Только админ может удалять пользователей
    if current_user.role_id != 1:  # Только администратор
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        user = get_object_or_404(Users, id_user=user_id)
        
        # Нельзя удалить самого себя
        if user.id_user == current_user_id:
            return JsonResponse({'error': 'Нельзя удалить свой собственный аккаунт'}, status=400)
        
        # Проверяем, есть ли у пользователя активные заказы или товары
        active_orders = Orders.objects.filter(user=user, status__in=['pending', 'paid']).exists()
        if active_orders:
            return JsonResponse({'error': 'Невозможно удалить пользователя с активными заказами'}, status=400)
        
        user_products = Products.objects.filter(seller=user, is_active=True).exists()
        if user_products:
            return JsonResponse({'error': 'Невозможно удалить пользователя с активными товарами'}, status=400)
        
        user_login = user.login
        user.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален пользователь: {user_login}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Пользователь успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Delete user error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def toggle_user_active(request, user_id):
    """Активация/деактивация пользователя (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    current_user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=current_user_id)
    
    # Только админ и менеджер могут менять статус
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        user = get_object_or_404(Users, id_user=user_id)
        
        # Нельзя деактивировать самого себя
        if user.id_user == current_user_id:
            return JsonResponse({'error': 'Нельзя деактивировать свой собственный аккаунт'}, status=400)
        
        # Менеджер не может деактивировать админов и других менеджеров
        if current_user.role_id == 3 and user.role_id in [1, 3]:
            return JsonResponse({'error': 'Менеджер не может деактивировать администраторов и других менеджеров'}, status=403)
        
        old_status = user.is_active
        user.is_active = not old_status
        user.save()
        
        action = "активирован" if user.is_active else "деактивирован"
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Пользователь {user.login} {action}'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Пользователь успешно {action}',
            'is_active': user.is_active
        })
        
    except Exception as e:
        logger.error(f'Toggle user active error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)




def admin_products(request):
    """Страница управления продуктами"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут управлять продуктами
    if current_user.role_id not in [1, 3]:  # Админ или менеджер
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    # Фильтрация и поиск
    search_query = request.GET.get('search', '')
    category_filter = request.GET.get('category', '')
    type_filter = request.GET.get('type', '')
    seller_filter = request.GET.get('seller', '')
    status_filter = request.GET.get('status', '')
    
    # Получаем все продукты с предварительной загрузкой связанных данных
    products = Products.objects.all().select_related(
        'seller', 'category', 'product_type'
    ).prefetch_related('tovars').order_by('-created_at')
    
    # Применяем фильтры
    if search_query:
        products = products.filter(
            Q(title__icontains=search_query) |
            Q(description__icontains=search_query)
        )
    
    if category_filter:
        products = products.filter(category_id=category_filter)
    
    if type_filter:
        products = products.filter(product_type_id=type_filter)
    
    if seller_filter:
        products = products.filter(seller_id=seller_filter)
    
    if status_filter:
        if status_filter == 'active':
            products = products.filter(is_active=True)
        elif status_filter == 'inactive':
            products = products.filter(is_active=False)
        elif status_filter == 'available':
            # Товары с доступными товарами (не проданными)
            products = products.annotate(
                available_count=Count('tovars', filter=Q(tovars__is_sold=False))
            ).filter(available_count__gt=0, is_active=True)
        elif status_filter == 'sold_out':
            # Товары без доступных товаров
            products = products.annotate(
                available_count=Count('tovars', filter=Q(tovars__is_sold=False))
            ).filter(available_count=0, is_active=True)
    
    # Получаем данные для фильтров
    categories = Categories.objects.all()
    product_types = ProductTypes.objects.all()
    sellers = Users.objects.filter(products__isnull=False).distinct()
    
    # Статистика
    stats = {
        'total_products': products.count(),
        'active_products': products.filter(is_active=True).count(),
        'total_tovars': Tovars.objects.count(),
        'available_tovars': Tovars.objects.filter(is_sold=False).count(),
        'total_revenue': OrderItems.objects.aggregate(
            total=Sum('price_at_time_of_purchase')
        )['total'] or 0,
    }
    
    # Пагинация
    paginator = Paginator(products, 15)  # 15 продуктов на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'admin/products.html', {
        'products': page_obj,
        'page_obj': page_obj,
        'categories': categories,
        'product_types': product_types,
        'sellers': sellers,
        'stats': stats,
        'search_query': search_query,
        'category_filter': category_filter,
        'type_filter': type_filter,
        'seller_filter': seller_filter,
        'status_filter': status_filter,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

def admin_product_tovars(request, product_id):
    """Получение товаров продукта (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        product = get_object_or_404(Products, id_product=product_id)
        
        # Получаем все товары продукта
        tovars = product.tovars.all().order_by('-created_at')
        
        tovars_data = []
        for tovar in tovars:
            tovars_data.append({
                'id': tovar.id_tovar,
                'text': tovar.tovar_text[:100] + ('...' if len(tovar.tovar_text) > 100 else ''),
                'full_text': tovar.tovar_text,
                'is_sold': tovar.is_sold,
                'sold_at': tovar.sold_at.strftime('%d.%m.%Y %H:%M') if tovar.sold_at else None,
                'created_at': tovar.created_at.strftime('%d.%m.%Y %H:%M')
            })
        
        return JsonResponse({
            'success': True,
            'product': {
                'id': product.id_product,
                'title': product.title,
                'seller': product.seller.login
            },
            'tovars': tovars_data,
            'total_tovars': tovars.count(),
            'available_tovars': product.available_tovars_count
        })
        
    except Exception as e:
        logger.error(f'Admin get product tovars error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_create_product(request):
    """Создание нового продукта (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут создавать продукты
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        # Получаем данные из POST (не из request.body, потому что это multipart)
        title = request.POST.get('title')
        description = request.POST.get('description')
        price = request.POST.get('price')
        seller_id = request.POST.get('seller_id')
        category_id = request.POST.get('category_id')
        product_type_id = request.POST.get('product_type_id')
        is_auto_delivery = request.POST.get('is_auto_delivery') == 'true'
        auto_delivery_text = request.POST.get('auto_delivery_text', '')
        
        # Получаем товары из поля tovars (это JSON строка)
        tovars_json = request.POST.get('tovars', '[]')
        tovars_data = json.loads(tovars_json)
        
        # Получаем загруженное изображение
        product_image = request.FILES.get('product_image')
        
        # Валидация
        if not all([title, description, price, seller_id, product_type_id]):
            return JsonResponse({'error': 'Заполните все обязательные поля'}, status=400)
        
        if not tovars_data:
            return JsonResponse({'error': 'Добавьте хотя бы один товар'}, status=400)
        
        try:
            # Преобразуем строку в Decimal
            price_decimal = Decimal(str(price).replace(',', '.'))
            if price_decimal <= 0:
                return JsonResponse({'error': 'Цена должна быть положительной'}, status=400)
        except (InvalidOperation, ValueError):
            return JsonResponse({'error': 'Некорректная цена'}, status=400)
        
        # Обработка изображения
        main_image_url = None
        if product_image:
            # Генерируем имя файла
            import os
            from django.utils.text import slugify
            import uuid
            
            # Создаем безопасное имя файла
            file_extension = os.path.splitext(product_image.name)[1]
            filename = f"{slugify(title)}_{uuid.uuid4().hex[:8]}{file_extension}"
            
            # Сохраняем файл
            from django.core.files.storage import default_storage
            from django.core.files.base import ContentFile
            
            file_path = default_storage.save(f'products/{filename}', ContentFile(product_image.read()))
            main_image_url = filename
        
        # Создаем продукт
        product = Products.objects.create(
            title=title,
            description=description,
            price=price_decimal,
            seller_id=seller_id,
            category_id=category_id if category_id else None,
            product_type_id=product_type_id,
            main_image_url=main_image_url,
            is_auto_delivery=bool(is_auto_delivery),
            auto_delivery_text=auto_delivery_text,
            is_active=True
        )
        
        # Добавляем товары
        for tovar_text in tovars_data:
            if tovar_text.strip():
                tovar = Tovars.objects.create(
                    tovar_text=tovar_text.strip(),
                    is_sold=False
                )
                ProductsTovars.objects.create(
                    product=product,
                    tovar=tovar
                )
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Создан новый продукт: {title} с {len(tovars_data)} товарами'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Продукт успешно создан',
            'product_id': product.id_product
        })
        
    except Exception as e:
        logger.error(f'Admin create product error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_edit_product(request, product_id):
    """Редактирование продукта (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут редактировать продукты
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        product = get_object_or_404(Products, id_product=product_id)
        
        # Получаем данные из POST
        title = request.POST.get('title')
        description = request.POST.get('description')
        price = request.POST.get('price')
        category_id = request.POST.get('category_id')
        product_type_id = request.POST.get('product_type_id')
        is_auto_delivery = request.POST.get('is_auto_delivery') == 'true'
        auto_delivery_text = request.POST.get('auto_delivery_text', '')
        
        # Получаем загруженное изображение
        product_image = request.FILES.get('product_image')
        
        # Валидация
        if not all([title, description, price, product_type_id]):
            return JsonResponse({'error': 'Заполните все обязательные поля'}, status=400)
        
        try:
            # Преобразуем строку в Decimal
            price_decimal = Decimal(str(price).replace(',', '.'))
            if price_decimal <= 0:
                return JsonResponse({'error': 'Цена должна быть положительной'}, status=400)
        except (InvalidOperation, ValueError):
            return JsonResponse({'error': 'Некорректная цена'}, status=400)
        
        old_title = product.title
        
        # Обработка изображения
        if product_image:
            # Удаляем старое изображение, если оно есть
            if product.main_image_url:
                try:
                    from django.core.files.storage import default_storage
                    old_file_path = f'products/{product.main_image_url}'
                    if default_storage.exists(old_file_path):
                        default_storage.delete(old_file_path)
                except:
                    pass
            
            # Генерируем имя файла
            import os
            from django.utils.text import slugify
            import uuid
            
            file_extension = os.path.splitext(product_image.name)[1]
            filename = f"{slugify(title)}_{uuid.uuid4().hex[:8]}{file_extension}"
            
            # Сохраняем файл
            from django.core.files.storage import default_storage
            from django.core.files.base import ContentFile
            
            file_path = default_storage.save(f'products/{filename}', ContentFile(product_image.read()))
            product.main_image_url = filename
        
        product.title = title
        product.description = description
        product.price = price_decimal
        product.category_id = category_id if category_id else None
        product.product_type_id = product_type_id
        product.is_auto_delivery = bool(is_auto_delivery)
        product.auto_delivery_text = auto_delivery_text
        product.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Отредактирован продукт: {old_title} -> {title}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Продукт успешно обновлен'
        })
        
    except Exception as e:
        logger.error(f'Admin edit product error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_delete_product(request, product_id):
    """Удаление продукта (API) - удаление только неактивных продуктов"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ может удалять продукты
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        product = get_object_or_404(Products, id_product=product_id)
        
        # Проверяем активность продукта
        if product.is_active:
            return JsonResponse({
                'error': 'Невозможно удалить активный продукт. Сначала деактивируйте продукт.'
            }, status=400)
        
        product_title = product.title
        
        # Удаляем все непроданные товары продукта
        tovars_to_delete = []
        sold_tovars_count = 0
        
        for tovar in product.tovars.all():
            if tovar.is_sold:
                sold_tovars_count += 1
                # Для проданных товаров разрываем связь, но не удаляем сам товар
                ProductsTovars.objects.filter(product=product, tovar=tovar).delete()
            else:
                # Проверяем, используется ли товар в других продуктах
                other_products_count = tovar.products.exclude(id_product=product_id).count()
                if other_products_count == 0:
                    tovars_to_delete.append(tovar)
                else:
                    # Если товар используется в других продуктах, только разрываем связь
                    ProductsTovars.objects.filter(product=product, tovar=tovar).delete()
        
        # Удаляем сам продукт
        product.delete()
        
        # Удаляем товары, которые больше нигде не используются
        deleted_tovars_count = 0
        for tovar in tovars_to_delete:
            tovar.delete()
            deleted_tovars_count += 1
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален неактивный продукт: {product_title} (удалено товаров: {deleted_tovars_count}, сохранено проданных: {sold_tovars_count})'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Продукт успешно удален. Удалено {deleted_tovars_count} товаров, сохранено {sold_tovars_count} проданных товаров.',
            'deleted_tovars': deleted_tovars_count,
            'saved_tovars': sold_tovars_count
        })
        
    except Exception as e:
        logger.error(f'Admin delete product error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)




def admin_orders(request):
    """Страница управления заказами"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут управлять заказами
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    # Фильтрация и поиск
    search_query = request.GET.get('search', '')
    status_filter = request.GET.get('status', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    user_filter = request.GET.get('user', '')
    
    # Получаем все заказы с предварительной загрузкой связанных данных
    orders = Orders.objects.all().select_related('user').prefetch_related(
        'orderitems_set__product'
    ).order_by('-order_created_at')
    
    # Применяем фильтры
    if search_query:
        orders = orders.filter(
            Q(id_order__icontains=search_query) |
            Q(user__login__icontains=search_query) |
            Q(user__firstname__icontains=search_query) |
            Q(user__surname__icontains=search_query) |
            Q(payment_reference__icontains=search_query)
        )
    
    if status_filter:
        orders = orders.filter(status=status_filter)
    
    if date_from:
        try:
            date_from_obj = datetime.strptime(date_from, '%Y-%m-%d')
            orders = orders.filter(order_created_at__date__gte=date_from_obj)
        except ValueError:
            pass
    
    if date_to:
        try:
            date_to_obj = datetime.strptime(date_to, '%Y-%m-%d')
            orders = orders.filter(order_created_at__date__lte=date_to_obj)
        except ValueError:
            pass
    
    if user_filter:
        orders = orders.filter(user_id=user_filter)
    
    # Получаем данные для фильтров
    users = Users.objects.filter(orders__isnull=False).distinct().order_by('login')
    
    # Статистика
    today = timezone.now().date()
    
    stats = {
        'total_orders': orders.count(),
        'pending_orders': orders.filter(status='pending').count(),
        'paid_orders': orders.filter(status='paid').count(),
        'completed_orders': orders.filter(status='completed').count(),
        'today_orders': orders.filter(order_created_at__date=today).count(),
        'today_revenue': orders.filter(
            order_created_at__date=today, 
            status__in=['paid', 'completed']
        ).aggregate(total=Sum('total_cost'))['total'] or 0,
        'total_revenue': orders.filter(
            status__in=['paid', 'completed']
        ).aggregate(total=Sum('total_cost'))['total'] or 0,
    }
    
    # Пагинация
    paginator = Paginator(orders, 15)  # 15 заказов на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'admin/orders.html', {
        'orders': page_obj,
        'page_obj': page_obj,
        'users': users,
        'stats': stats,
        'search_query': search_query,
        'status_filter': status_filter,
        'date_from': date_from,
        'date_to': date_to,
        'user_filter': user_filter,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

def admin_order_detail(request, order_id):
    """Детальная информация о заказе"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут просматривать заказы
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    try:
        order = Orders.objects.select_related('user').get(id_order=order_id)
        
        # Получаем все позиции заказа с информацией о товарах
        order_items = OrderItems.objects.filter(
            order=order
        ).select_related('product', 'tovar')
        
        # Получаем транзакции, связанные с заказом
        transactions = Transactions.objects.filter(order=order).order_by('-created_at')
        
        # Получаем чат, связанный с заказом (если есть)
        chats = Chats.objects.filter(order=order).select_related('buyer', 'seller')
        
        return render(request, 'admin/order_detail.html', {
            'order': order,
            'order_items': order_items,
            'transactions': transactions,
            'chats': chats,
            'is_admin': current_user.role_id == 1,
            'is_manager': current_user.role_id == 3
        })
        
    except Orders.DoesNotExist:
        messages.error(request, 'Заказ не найден')
        return redirect('admin_orders')


# appip/views.py - СТАБИЛЬНАЯ ВЕРСИЯ + СПИСАНИЕ С ПРОДАВЦОВ

@csrf_exempt
@require_POST
def admin_update_order_status(request, order_id):
    """Обновление статуса заказа (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут обновлять статусы заказов
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        order = get_object_or_404(Orders, id_order=order_id)
        
        data = json.loads(request.body)
        new_status = data.get('status')
        
        # Валидация статуса
        valid_statuses = ['pending', 'completed', 'cancelled']
        if new_status not in valid_statuses:
            return JsonResponse({'error': 'Неверный статус'}, status=400)
        
        old_status = order.status
        
        # Проверяем переходы статусов
        if old_status == 'cancelled' and new_status != 'cancelled':
            return JsonResponse({'error': 'Невозможно изменить статус отмененного заказа'}, status=400)
        
        # Если заказ отменен, возвращаем товары на склад и средства
        if new_status == 'cancelled' and old_status != 'cancelled':
            # Получаем все товары в заказе
            order_items = OrderItems.objects.filter(order=order).select_related('tovar', 'product__seller')
            
            # 1. Возвращаем товары на склад
            for item in order_items:
                if item.tovar:
                    item.tovar.is_sold = False
                    item.tovar.sold_at = None
                    item.tovar.save()
            
            # 2. Если заказ был завершен, списываем деньги с продавцов
            if old_status == 'completed':
                for item in order_items:
                    seller = item.product.seller
                    amount = item.price_at_time_of_purchase * item.quantity
                    
                    if seller.balance >= amount:
                        seller.balance -= amount
                        seller.save()
                        
                        Transactions.objects.create(
                            user=seller,
                            order=order,
                            amount=amount,
                            transaction_type='refund',
                            status='completed',
                            reference=f'CANCEL_ORDER_{order.id_order}_SELLER_{seller.id_user}'
                        )
                        
                        UserActivityLog.objects.create(
                            user=seller,
                            action='admin_action',
                            description=f'Списание средств за отмененный заказ №{order.id_order}: -{amount} ₽'
                        )
            
            # 3. Возвращаем средства покупателю
            user = order.user
            user.balance += order.total_cost
            user.save()
            
            Transactions.objects.create(
                user=user,
                order=order,
                amount=order.total_cost,
                transaction_type='refund',
                status='completed',
                reference=f'CANCEL_ORDER_{order.id_order}'
            )
            
            UserActivityLog.objects.create(
                user=user,
                action='admin_action',
                description=f'Возврат средств за отмененный заказ №{order.id_order}: +{order.total_cost} ₽'
            )
        
        # Если заказ завершен, начисляем деньги продавцам
        elif new_status == 'completed' and old_status == 'pending':
            order_items = OrderItems.objects.filter(order=order).select_related('product__seller')
            
            for item in order_items:
                seller = item.product.seller
                amount = item.price_at_time_of_purchase * item.quantity
                
                seller.balance += amount
                seller.save()
                
                Transactions.objects.create(
                    user=seller,
                    order=order,
                    amount=amount,
                    transaction_type='sale',
                    status='completed',
                    reference=f'SALE_ORDER_{order.id_order}_SELLER_{seller.id_user}'
                )
                
                UserActivityLog.objects.create(
                    user=seller,
                    action='create_order',
                    description=f'Зачисление средств за заказ №{order.id_order}: +{amount} ₽'
                )
            
            order_items.update(status='delivered')
        
        # Меняем статус заказа
        order.status = new_status
        order.save()
        
        # Логирование действия администратора
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Изменен статус заказа #{order.id_order}: {old_status} -> {new_status}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Статус заказа успешно обновлен',
            'new_status': new_status,
            'status_display': order.get_status_display()
        })
        
    except Exception as e:
        logger.error(f'Admin update order status error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_delete_order(request, order_id):
    """Удаление заказа (API) - удаление только отмененных или возвращенных заказов"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    try:
        order = get_object_or_404(Orders, id_order=order_id)
        
        # Проверяем, можно ли удалить заказ
      
        if order.status not in ['cancelled', 'refunded']:
            return JsonResponse({
                'error': 'Невозможно удалить заказ. Заказ должен быть отменен или возвращен.'
            }, status=400)
        
        order_number = order.id_order
        
        # Логирование перед удалением
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален заказ #{order_number} (статус: {order.status})'
        )
        
        # Удаляем заказ
        order.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Заказ #{order_number} успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Admin delete order error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)








# ==================== УПРАВЛЕНИЕ ПРОМОКОДАМИ ====================
from django.utils import timezone


def admin_promocodes(request):
    """Страница управления промокодами"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    # Только админ может управлять промокодами
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    promocodes = PromoCodes.objects.all().order_by('-created_at')
    
    # Добавляем текущее время в контекст
    now = timezone.now()
    
    return render(request, 'admin/promocodes.html', {
        'promocodes': promocodes,
        'is_admin': current_user.role_id == 1,
        'current_user': current_user,
        'now': now 
    })

@csrf_exempt
@require_POST
def create_promocode(request):
    """Создание нового промокода"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        code = data.get('code', '').upper().strip()
        discount_percent = int(data.get('discount_percent', 0))
        expires_at = data.get('expires_at')
        usage_limit = int(data.get('usage_limit', 0))
        
        # Валидация
        if not code:
            return JsonResponse({'error': 'Введите код промокода'}, status=400)
        
        if PromoCodes.objects.filter(code=code).exists():
            return JsonResponse({'error': 'Промокод с таким именем уже существует'}, status=400)
        
        if discount_percent < 0 or discount_percent > 100:
            return JsonResponse({'error': 'Скидка должна быть от 0 до 100%'}, status=400)
        
        # Создаем промокод
        promocode = PromoCodes.objects.create(
            code=code,
            discount_percent=discount_percent,
            is_active=True,
            expires_at=expires_at if expires_at else None,
            usage_limit=usage_limit
        )
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Создан промокод: {code} ({discount_percent}%)'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Промокод успешно создан',
            'promocode': {
                'id': promocode.id_promocode,
                'code': promocode.code,
                'discount_percent': promocode.discount_percent,
                'is_active': promocode.is_active,
                'status_display': 'Активен'
            }
        })
        
    except Exception as e:
        logger.error(f'Create promocode error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def edit_promocode(request, promocode_id):
    """Редактирование промокода"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        promocode = get_object_or_404(PromoCodes, id_promocode=promocode_id)
        data = json.loads(request.body)
        
        code = data.get('code', '').upper().strip()
        discount_percent = int(data.get('discount_percent', 0))
        expires_at = data.get('expires_at')
        usage_limit = int(data.get('usage_limit', 0))
        
        # Валидация
        if not code:
            return JsonResponse({'error': 'Введите код промокода'}, status=400)
        
        if PromoCodes.objects.filter(code=code).exclude(id_promocode=promocode_id).exists():
            return JsonResponse({'error': 'Промокод с таким именем уже существует'}, status=400)
        
        if discount_percent < 0 or discount_percent > 100:
            return JsonResponse({'error': 'Скидка должна быть от 0 до 100%'}, status=400)
        
        old_code = promocode.code
        promocode.code = code
        promocode.discount_percent = discount_percent
        promocode.expires_at = expires_at if expires_at else None
        promocode.usage_limit = usage_limit
        promocode.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Отредактирован промокод: {old_code} -> {code} ({discount_percent}%)'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Промокод успешно обновлен'
        })
        
    except Exception as e:
        logger.error(f'Edit promocode error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def delete_promocode(request, promocode_id):
    """Удаление промокода"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        promocode = get_object_or_404(PromoCodes, id_promocode=promocode_id)
        code = promocode.code
        promocode.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален промокод: {code}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Промокод успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Delete promocode error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)



@csrf_exempt
@require_POST
def toggle_promocode(request, promocode_id):
    """Активация/деактивация промокода"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        promocode = get_object_or_404(PromoCodes, id_promocode=promocode_id)
        
        # Проверяем, не пытается ли пользователь повторно нажать
        old_status = promocode.is_active
        new_status = not old_status
        
        # Если статус уже изменен, не делаем ничего
        if promocode.is_active == new_status:
            return JsonResponse({
                'success': False,
                'error': 'Статус уже изменен'
            }, status=400)
        
        promocode.is_active = new_status
        promocode.save()
        
        action = "активирован" if promocode.is_active else "деактивирован"
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Промокод {promocode.code} {action}'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Промокод успешно {action}',
            'is_active': promocode.is_active
        })
        
    except Exception as e:
        logger.error(f'Toggle promocode error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


from decimal import Decimal



@csrf_exempt
@require_POST
def apply_promocode(request):
    """Применение промокода к корзине или проверка валидности"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        data = json.loads(request.body)
        code = data.get('code', '').upper().strip()
        validate_only = data.get('validate_only', False)
        
        if not code:
            return JsonResponse({'error': 'Введите код промокода'}, status=400)
        
       
        try:
            promocode = PromoCodes.objects.get(code=code)
        except PromoCodes.DoesNotExist:
            return JsonResponse({'error': 'Промокод не найден'}, status=404)
        
        if not promocode.is_valid:
            return JsonResponse({'error': 'Промокод недействителен'}, status=400)
        
        # Если только проверка, возвращаем успех без сохранения в сессию
        if validate_only:
            return JsonResponse({
                'success': True,
                'message': 'Промокод действителен'
            })
        
        cart_items = Cart.objects.filter(user_id=user_id).select_related('product')
        
        if not cart_items.exists():
            return JsonResponse({'error': 'Корзина пуста'}, status=400)
        
        total_price = Decimal('0')
        for item in cart_items:
            total_price += item.product.price * Decimal(str(item.quantity))
        
        discounted_price = promocode.apply_discount(total_price)
        discount_amount = total_price - discounted_price
        
       
        request.session['applied_promocode'] = {
            'code': promocode.code,
            'discount_percent': promocode.discount_percent,
            'original_price': float(total_price),
            'final_price': float(discounted_price)
        }
        
        return JsonResponse({
            'success': True,
            'message': f'Промокод применен! Скидка {promocode.discount_percent}%',
            'total_price': float(total_price),
            'discount_amount': float(discount_amount),
            'final_price': float(discounted_price),
            'discount_percent': promocode.discount_percent,
            'code': promocode.code
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный формат запроса'}, status=400)
    except Exception as e:
        logger.error(f'Apply promocode error: {str(e)}')
        return JsonResponse({'error': 'Произошла ошибка при применении промокода'}, status=500)



@csrf_exempt
@require_POST
def clear_promocode(request):
    """Очистка промокода из сессии"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    try:
        # Удаляем промокод из сессии
        request.session.pop('applied_promocode', None)
        return JsonResponse({'success': True})
    except Exception as e:
        logger.error(f'Clear promocode error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


import os
import json
import shutil
import zipfile
from datetime import datetime
from django.conf import settings
from django.http import FileResponse
from django.views.decorators.http import require_GET
from django.core import serializers
from django.db import connection
from django.db.models import Q
import tempfile

# ==================== ВОССТАНОВЛЕНИЕ ДАННЫХ ====================

def admin_data_recovery(request):
    """Страница восстановления данных и бэкапов"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    # Только админ может управлять бэкапами
    if current_user.role_id != 1:  # Только администратор
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    # Получаем список бэкапов
    backup_dir = os.path.join(settings.BASE_DIR, 'backups')
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)
    
    backups = []
    for filename in os.listdir(backup_dir):
        if filename.endswith('.zip'):
            filepath = os.path.join(backup_dir, filename)
            stat = os.stat(filepath)
            
            # Определяем размер в удобном формате
            size_bytes = stat.st_size
            if size_bytes < 1024:
                size_str = f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                size_str = f"{size_bytes / 1024:.1f} KB"
            else:
                size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
            
            backups.append({
                'filename': filename,
                'size': size_str,
                'created_at': datetime.fromtimestamp(stat.st_ctime).strftime('%d.%m.%Y %H:%M:%S'),
                'modified_at': datetime.fromtimestamp(stat.st_mtime).strftime('%d.%m.%Y %H:%M:%S')
            })
    
    # Сортировка по дате создания (новые сверху)
    backups.sort(key=lambda x: x['created_at'], reverse=True)
    
    # Статистика базы данных
    db_stats = {}
    try:
        with connection.cursor() as cursor:
            # Получаем размер базы данных
            cursor.execute("SELECT pg_database_size(current_database())")
            db_stats['size'] = cursor.fetchone()[0]
            
            # Получаем количество записей в основных таблицах
            tables = [
                'users', 'products', 'tovars', 'orders', 'orderitems',
                'categories', 'producttypes', 'productreviews', 'sellerreviews',
                'wishlists', 'cart', 'chats', 'messages', 'transactions',
                'useractivitylogs'
            ]
            
            table_counts = {}
            for table in tables:
                try:
                    cursor.execute(f'SELECT COUNT(*) FROM "{table}"')
                    table_counts[table] = cursor.fetchone()[0]
                except:
                    table_counts[table] = 0
            
            db_stats['table_counts'] = table_counts
            
    except Exception as e:
        logger.error(f'Error getting DB stats: {str(e)}')
        db_stats['error'] = str(e)
    
    return render(request, 'admin/datarecovery.html', {
        'backups': backups,
        'db_stats': db_stats,
        'is_admin': current_user.role_id == 1
    })

@csrf_exempt
@require_POST
def admin_create_backup(request):
    """Создание резервной копии данных"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        data = json.loads(request.body)
        backup_type = data.get('backup_type', 'full')  # full, structure, data
        include_files = data.get('include_files', False)
        
        # Создаем директорию для бэкапов
        backup_dir = os.path.join(settings.BASE_DIR, 'backups')
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
        
        # Имя файла с датой
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_filename = f"backup_{timestamp}_{backup_type}.zip"
        backup_path = os.path.join(backup_dir, backup_filename)
        
        # Создаем временную директорию
        with tempfile.TemporaryDirectory() as temp_dir:
            # Экспортируем данные из моделей
            models_to_export = {
                'roles': Roles,
                'users': Users,
                'categories': Categories,
                'producttypes': ProductTypes,
                'tovars': Tovars,
                'products': Products,
                'productstovars': ProductsTovars,
                'productcategories': ProductCategories,
                'productitems': ProductItems,
                'orders': Orders,
                'orderitems': OrderItems,
                'productreviews': ProductReviews,
                'sellerreviews': SellerReviews,
                'wishlists': Wishlists,
                'cart': Cart,
                'chats': Chats,
                'messages': Messages,
                'transactions': Transactions,
                'useractivitylogs': UserActivityLog,
            }
            
            if include_files:
                models_to_export['managerchats'] = ManagerChats
            
            # Экспортируем каждую модель в JSON
            for name, model in models_to_export.items():
                try:
                    data = serializers.serialize('json', model.objects.all())
                    with open(os.path.join(temp_dir, f'{name}.json'), 'w', encoding='utf-8') as f:
                        f.write(data)
                except Exception as e:
                    logger.error(f'Error exporting {name}: {str(e)}')
            
            # Экспортируем структуру базы данных (схему)
            if backup_type in ['full', 'structure']:
                with connection.cursor() as cursor:
                    cursor.execute("""
                        SELECT 
                            table_name, 
                            column_name, 
                            data_type, 
                            is_nullable,
                            column_default
                        FROM information_schema.columns 
                        WHERE table_schema = 'public'
                        ORDER BY table_name, ordinal_position
                    """)
                    schema = cursor.fetchall()
                    
                    schema_data = {}
                    for table, column, data_type, nullable, default in schema:
                        if table not in schema_data:
                            schema_data[table] = []
                        schema_data[table].append({
                            'column': column,
                            'type': data_type,
                            'nullable': nullable,
                            'default': default
                        })
                    
                    with open(os.path.join(temp_dir, 'schema.json'), 'w', encoding='utf-8') as f:
                        json.dump(schema_data, f, indent=2, ensure_ascii=False)
            
            # Создаем ZIP архив
            with zipfile.ZipFile(backup_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, temp_dir)
                        zipf.write(file_path, arcname)
            
            # Получаем размер файла
            file_size = os.path.getsize(backup_path)
            
            # Логирование
            UserActivityLog.objects.create(
                user=current_user,
                action='admin_action',
                description=f'Создан бэкап: {backup_filename} (тип: {backup_type}, размер: {file_size} байт)'
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Резервная копия успешно создана',
                'filename': backup_filename,
                'size': file_size,
                'type': backup_type
            })
            
    except Exception as e:
        logger.error(f'Create backup error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@require_GET
def admin_download_backup(request, filename):
    """Скачивание файла бэкапа"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    # Проверка безопасности пути
    if '..' in filename or filename.startswith('/'):
        return JsonResponse({'error': 'Некорректное имя файла'}, status=400)
    
    backup_dir = os.path.join(settings.BASE_DIR, 'backups')
    filepath = os.path.join(backup_dir, filename)
    
    if not os.path.exists(filepath):
        return JsonResponse({'error': 'Файл не найден'}, status=404)
    
    # Логирование
    UserActivityLog.objects.create(
        user=current_user,
        action='admin_action',
        description=f'Скачан бэкап: {filename}'
    )
    
    return FileResponse(
        open(filepath, 'rb'),
        as_attachment=True,
        filename=filename
    )

@csrf_exempt
@require_POST
def admin_restore_backup(request):
    """Восстановление данных из бэкапа"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        data = json.loads(request.body)
        filename = data.get('filename')
        restore_type = data.get('restore_type', 'data')  # data, structure, full
        confirm = data.get('confirm', False)
        
        if not confirm:
            return JsonResponse({
                'error': 'Необходимо подтверждение для восстановления данных'
            }, status=400)
        
        if '..' in filename or filename.startswith('/'):
            return JsonResponse({'error': 'Некорректное имя файла'}, status=400)
        
        backup_dir = os.path.join(settings.BASE_DIR, 'backups')
        filepath = os.path.join(backup_dir, filename)
        
        if not os.path.exists(filepath):
            return JsonResponse({'error': 'Файл бэкапа не найден'}, status=404)
        
        # Создаем временную директорию для распаковки
        with tempfile.TemporaryDirectory() as temp_dir:
            # Распаковываем архив
            with zipfile.ZipFile(filepath, 'r') as zipf:
                zipf.extractall(temp_dir)
            
            restored_models = []
            errors = []
            
            # Определяем порядок восстановления (учитывая внешние ключи)
            restore_order = [
                'roles.json',
                'users.json',
                'categories.json',
                'producttypes.json',
                'tovars.json',
                'products.json',
                'productstovars.json',
                'productcategories.json',
                'productitems.json',
                'orders.json',
                'orderitems.json',
                'productreviews.json',
                'sellerreviews.json',
                'wishlists.json',
                'cart.json',
                'chats.json',
                'messages.json',
                'transactions.json',
                'useractivitylogs.json',
                'managerchats.json'
            ]
            
            # Восстанавливаем данные
            for json_file in restore_order:
                json_path = os.path.join(temp_dir, json_file)
                if os.path.exists(json_path):
                    try:
                        with open(json_path, 'r', encoding='utf-8') as f:
                            data = f.read()
                        
                        # Десериализуем данные
                        objects = serializers.deserialize('json', data)
                        
                        model_name = json_file.replace('.json', '')
                        count = 0
                        
                        for obj in objects:
                            obj.save()
                            count += 1
                        
                        restored_models.append(f"{model_name}: {count} записей")
                        
                    except Exception as e:
                        errors.append(f"Ошибка восстановления {json_file}: {str(e)}")
            
            # Логирование
            UserActivityLog.objects.create(
                user=current_user,
                action='admin_action',
                description=f'Восстановление из бэкапа: {filename} (восстановлено: {", ".join(restored_models)})'
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Данные успешно восстановлены',
                'restored_models': restored_models,
                'errors': errors if errors else None
            })
            
    except Exception as e:
        logger.error(f'Restore backup error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_delete_backup(request, filename):
    """Удаление файла бэкапа"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        if '..' in filename or filename.startswith('/'):
            return JsonResponse({'error': 'Некорректное имя файла'}, status=400)
        
        backup_dir = os.path.join(settings.BASE_DIR, 'backups')
        filepath = os.path.join(backup_dir, filename)
        
        if not os.path.exists(filepath):
            return JsonResponse({'error': 'Файл не найден'}, status=404)
        
        # Удаляем файл
        os.remove(filepath)
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален бэкап: {filename}'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Бэкап {filename} успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Delete backup error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_upload_backup(request):
    """Загрузка файла бэкапа на сервер"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = get_object_or_404(Users, id_user=user_id)
    
    if current_user.role_id != 1:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        if 'backup_file' not in request.FILES:
            return JsonResponse({'error': 'Файл не выбран'}, status=400)
        
        uploaded_file = request.FILES['backup_file']
        
        # Проверка расширения
        if not uploaded_file.name.endswith('.zip'):
            return JsonResponse({'error': 'Файл должен быть в формате ZIP'}, status=400)
        
        # Проверка размера (макс 100MB)
        if uploaded_file.size > 100 * 1024 * 1024:
            return JsonResponse({'error': 'Размер файла превышает 100MB'}, status=400)
        
        # Создаем директорию для бэкапов
        backup_dir = os.path.join(settings.BASE_DIR, 'backups')
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
        
        # Сохраняем файл
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"uploaded_{timestamp}_{uploaded_file.name}"
        filepath = os.path.join(backup_dir, filename)
        
        with open(filepath, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Загружен бэкап: {filename}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Файл успешно загружен',
            'filename': filename
        })
        
    except Exception as e:
        logger.error(f'Upload backup error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)








@csrf_exempt
@require_POST
def admin_toggle_product_active(request, product_id):
    """Активация/деактивация продукта (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут менять статус
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        product = get_object_or_404(Products, id_product=product_id)
        
        old_status = product.is_active
        product.is_active = not old_status
        product.save()
        
        action = "активирован" if product.is_active else "деактивирован"
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Продукт {product.title} {action}'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Продукт успешно {action}',
            'is_active': product.is_active
        })
        
    except Exception as e:
        logger.error(f'Admin toggle product active error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_delete_tovar(request, tovar_id):
    """Удаление товара (Tovar) (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут удалять товары
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        tovar = get_object_or_404(Tovars, id_tovar=tovar_id)
        
        # Проверяем, продан ли товар
        if tovar.is_sold:
            return JsonResponse({'error': 'Невозможно удалить проданный товар'}, status=400)
        
        # Получаем продукты, связанные с этим товаром
        products = tovar.products.all()
        products_info = [{'id': p.id_product, 'title': p.title} for p in products]
        
        # Удаляем связи из ProductsTovars
        ProductsTovars.objects.filter(tovar=tovar).delete()
        
        # Удаляем сам товар
        tovar.delete()
        
        # Логирование
        products_titles = ', '.join([p['title'] for p in products_info])
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален товар #{tovar_id} из продуктов: {products_titles}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Товар успешно удален',
            'affected_products': products_info
        })
        
    except Exception as e:
        logger.error(f'Admin delete tovar error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)






def admin_wishlists(request):
    """Страница управления вишлистами"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут управлять вишлистами
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    # Фильтрация и поиск
    search_query = request.GET.get('search', '')
    user_filter = request.GET.get('user', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    
    # Получаем все вишлисты с предварительной загрузкой связанных данных
    wishlists = Wishlists.objects.all().select_related('user', 'product').order_by('-added_at')
    
    # Применяем фильтры
    if search_query:
        wishlists = wishlists.filter(
            Q(user__login__icontains=search_query) |
            Q(user__firstname__icontains=search_query) |
            Q(user__surname__icontains=search_query) |
            Q(product__title__icontains=search_query) |
            Q(product__description__icontains=search_query)
        )
    
    if user_filter:
        wishlists = wishlists.filter(user_id=user_filter)
    
    if date_from:
        try:
            date_from_obj = datetime.strptime(date_from, '%Y-%m-%d')
            wishlists = wishlists.filter(added_at__date__gte=date_from_obj)
        except ValueError:
            pass
    
    if date_to:
        try:
            date_to_obj = datetime.strptime(date_to, '%Y-%m-%d')
            wishlists = wishlists.filter(added_at__date__lte=date_to_obj)
        except ValueError:
            pass
    
    # Получаем данные для фильтров
    users = Users.objects.filter(wishlists__isnull=False).distinct().order_by('login')
    
    # Статистика
    today = timezone.now().date()
    
    stats = {
        'total_wishlists': wishlists.count(),
        'unique_users': wishlists.values('user').distinct().count(),
        'unique_products': wishlists.values('product').distinct().count(),
        'today_added': wishlists.filter(added_at__date=today).count(),
        'most_popular_product': wishlists.values('product__title').annotate(
            count=Count('id_wishlist')
        ).order_by('-count').first(),
    }
    
    # Пагинация
    paginator = Paginator(wishlists, 15)  # 15 записей на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'admin/wishlists.html', {
        'wishlists': page_obj,
        'page_obj': page_obj,
        'users': users,
        'stats': stats,
        'search_query': search_query,
        'user_filter': user_filter,
        'date_from': date_from,
        'date_to': date_to,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

@csrf_exempt
@require_POST
def admin_delete_wishlist(request, wishlist_id):
    """Удаление записи из вишлиста (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут удалять записи вишлистов
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        wishlist = get_object_or_404(Wishlists, id_wishlist=wishlist_id)
        
        user_info = f"{wishlist.user.login} ({wishlist.user.firstname} {wishlist.user.surname})"
        product_info = wishlist.product.title
        added_date = wishlist.added_at.strftime('%d.%m.%Y %H:%M')
        
        # Удаляем запись
        wishlist.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удалена запись из вишлиста: {user_info} → {product_info} (добавлено: {added_date})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Запись успешно удалена из вишлиста'
        })
        
    except Exception as e:
        logger.error(f'Admin delete wishlist error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)





def admin_carts(request):
    """Страница управления корзинами"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут управлять корзинами
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    # Фильтрация и поиск
    search_query = request.GET.get('search', '')
    user_filter = request.GET.get('user', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    quantity_min = request.GET.get('quantity_min', '')
    quantity_max = request.GET.get('quantity_max', '')
    
    # Получаем все корзины с предварительной загрузкой связанных данных
    carts = Cart.objects.all().select_related('user', 'product').order_by('-added_at')
    
    # Применяем фильтры
    if search_query:
        carts = carts.filter(
            Q(user__login__icontains=search_query) |
            Q(user__firstname__icontains=search_query) |
            Q(user__surname__icontains=search_query) |
            Q(product__title__icontains=search_query) |
            Q(product__description__icontains=search_query)
        )
    
    if user_filter:
        carts = carts.filter(user_id=user_filter)
    
    if date_from:
        try:
            date_from_obj = datetime.strptime(date_from, '%Y-%m-%d')
            carts = carts.filter(added_at__date__gte=date_from_obj)
        except ValueError:
            pass
    
    if date_to:
        try:
            date_to_obj = datetime.strptime(date_to, '%Y-%m-%d')
            carts = carts.filter(added_at__date__lte=date_to_obj)
        except ValueError:
            pass
    
    if quantity_min:
        try:
            carts = carts.filter(quantity__gte=int(quantity_min))
        except ValueError:
            pass
    
    if quantity_max:
        try:
            carts = carts.filter(quantity__lte=int(quantity_max))
        except ValueError:
            pass
    
    # Получаем данные для фильтров
    users = Users.objects.filter(cart__isnull=False).distinct().order_by('login')
    
    # Статистика
    today = timezone.now().date()
    
    # Вычисляем общую стоимость товаров в корзинах
    total_value = 0
    for cart in carts:
        total_value += cart.product.price * cart.quantity
    
    stats = {
        'total_carts': carts.count(),
        'unique_users': carts.values('user').distinct().count(),
        'unique_products': carts.values('product').distinct().count(),
        'today_added': carts.filter(added_at__date=today).count(),
        'total_items': carts.aggregate(total=Sum('quantity'))['total'] or 0,
        'total_value': total_value,
        'avg_quantity': carts.aggregate(avg=Avg('quantity'))['avg'] or 0,
        'most_popular_product': carts.values('product__title').annotate(
            total_quantity=Sum('quantity'),
            count=Count('id_cart')
        ).order_by('-total_quantity').first(),
    }
    
    # Пагинация
    paginator = Paginator(carts, 15)  # 15 записей на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'admin/carts.html', {
        'carts': page_obj,
        'page_obj': page_obj,
        'users': users,
        'stats': stats,
        'search_query': search_query,
        'user_filter': user_filter,
        'date_from': date_from,
        'date_to': date_to,
        'quantity_min': quantity_min,
        'quantity_max': quantity_max,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

@csrf_exempt
@require_POST
def admin_delete_cart(request, cart_id):
    """Удаление записи из корзины (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут удалять записи корзин
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        cart = get_object_or_404(Cart, id_cart=cart_id)
        
        user_info = f"{cart.user.login} ({cart.user.firstname} {cart.user.surname})"
        product_info = cart.product.title
        quantity = cart.quantity
        total_price = cart.product.price * quantity
        added_date = cart.added_at.strftime('%d.%m.%Y %H:%M')
        
        # Удаляем запись
        cart.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удалена запись из корзины: {user_info} → {product_info} (x{quantity}, {total_price}₽, добавлено: {added_date})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Товар успешно удален из корзины'
        })
        
    except Exception as e:
        logger.error(f'Admin delete cart error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)





def admin_reviews(request):
    """Страница управления отзывами"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут управлять отзывами
    if current_user.role_id not in [1, 3]:
        messages.error(request, 'Недостаточно прав')
        return redirect('admin_dashboard')
    
    # Фильтрация и поиск для отзывов на товары
    product_search = request.GET.get('product_search', '')
    product_user_filter = request.GET.get('product_user', '')
    product_rating_filter = request.GET.get('product_rating', '')
    
    # Получаем отзывы на товары с предварительной загрузкой данных
    product_reviews = ProductReviews.objects.all().select_related(
        'user', 'product'
    ).order_by('-created_at')
    
    # Применяем фильтры для отзывов на товары
    if product_search:
        product_reviews = product_reviews.filter(
            Q(product__title__icontains=product_search) |
            Q(review_text__icontains=product_search) |
            Q(user__login__icontains=product_search) |
            Q(user__firstname__icontains=product_search) |
            Q(user__surname__icontains=product_search)
        )
    
    if product_user_filter:
        product_reviews = product_reviews.filter(user_id=product_user_filter)
    
    if product_rating_filter:
        product_reviews = product_reviews.filter(rating=product_rating_filter)
    
    # Фильтрация и поиск для отзывов на продавцов
    seller_search = request.GET.get('seller_search', '')
    seller_buyer_filter = request.GET.get('seller_buyer', '')
    seller_seller_filter = request.GET.get('seller_seller', '')
    seller_rating_filter = request.GET.get('seller_rating', '')
    
    # Получаем отзывы на продавцов с предварительной загрузкой данных
    seller_reviews = SellerReviews.objects.all().select_related(
        'buyer', 'seller'
    ).order_by('-created_at')
    
    # Применяем фильтры для отзывов на продавцов
    if seller_search:
        seller_reviews = seller_reviews.filter(
            Q(review_text__icontains=seller_search) |
            Q(seller__login__icontains=seller_search) |
            Q(seller__firstname__icontains=seller_search) |
            Q(seller__surname__icontains=seller_search) |
            Q(buyer__login__icontains=seller_search) |
            Q(buyer__firstname__icontains=seller_search) |
            Q(buyer__surname__icontains=seller_search)
        )
    
    if seller_buyer_filter:
        seller_reviews = seller_reviews.filter(buyer_id=seller_buyer_filter)
    
    if seller_seller_filter:
        seller_reviews = seller_reviews.filter(seller_id=seller_seller_filter)
    
    if seller_rating_filter:
        seller_reviews = seller_reviews.filter(rating=seller_rating_filter)
    
    # Получаем данные для фильтров
    product_users = Users.objects.filter(productreviews__isnull=False).distinct().order_by('login')
    seller_buyers = Users.objects.filter(buyer_reviews__isnull=False).distinct().order_by('login')
    seller_sellers = Users.objects.filter(seller_reviews__isnull=False).distinct().order_by('login')
    
    # Статистика
    today = timezone.now().date()
    
    stats = {
        'total_product_reviews': product_reviews.count(),
        'total_seller_reviews': seller_reviews.count(),
        'avg_product_rating': product_reviews.aggregate(avg=Avg('rating'))['avg'] or 0,
        'avg_seller_rating': seller_reviews.aggregate(avg=Avg('rating'))['avg'] or 0,
        'today_product_reviews': product_reviews.filter(created_at__date=today).count(),
        'today_seller_reviews': seller_reviews.filter(created_at__date=today).count(),
        'edited_reviews': product_reviews.filter(is_edited=True).count(),
    }
    
    # Пагинация
    product_paginator = Paginator(product_reviews, 10)  # 10 отзывов на товары на страницу
    seller_paginator = Paginator(seller_reviews, 10)    # 10 отзывов на продавцов на страницу
    
    product_page = request.GET.get('product_page', 1)
    seller_page = request.GET.get('seller_page', 1)
    
    product_page_obj = product_paginator.get_page(product_page)
    seller_page_obj = seller_paginator.get_page(seller_page)
    
    return render(request, 'admin/reviews.html', {
        'product_reviews': product_page_obj,
        'seller_reviews': seller_page_obj,
        'product_page_obj': product_page_obj,
        'seller_page_obj': seller_page_obj,
        'product_users': product_users,
        'seller_buyers': seller_buyers,
        'seller_sellers': seller_sellers,
        'stats': stats,
        'product_search': product_search,
        'product_user_filter': product_user_filter,
        'product_rating_filter': product_rating_filter,
        'seller_search': seller_search,
        'seller_buyer_filter': seller_buyer_filter,
        'seller_seller_filter': seller_seller_filter,
        'seller_rating_filter': seller_rating_filter,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

@csrf_exempt
@require_POST
def admin_delete_product_review(request, review_id):
    """Удаление отзыва на товар (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут удалять отзывы
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        review = get_object_or_404(ProductReviews, id_review=review_id)
        
        user_info = f"{review.user.login} ({review.user.firstname} {review.user.surname})"
        product_info = review.product.title
        rating = review.rating
        created_date = review.created_at.strftime('%d.%m.%Y %H:%M')
        
        # Удаляем отзыв
        review.delete()
        
        
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален отзыв на товар: {user_info} → {product_info} (рейтинг: {rating}, дата: {created_date})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Отзыв на товар успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Admin delete product review error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_delete_seller_review(request, review_id):
    """Удаление отзыва на продавца (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ и менеджер могут удалять отзывы
    if current_user.role_id not in [1, 3]:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        review = get_object_or_404(SellerReviews, id_sellerreview=review_id)
        
        buyer_info = f"{review.buyer.login} ({review.buyer.firstname} {review.buyer.surname})"
        seller_info = f"{review.seller.login} ({review.seller.firstname} {review.seller.surname})"
        rating = review.rating
        created_date = review.created_at.strftime('%d.%m.%Y %H:%M')
        
        # Удаляем отзыв
        review.delete()
        
        # Пересчитываем рейтинг продавца
        update_seller_rating(review.seller_id)
        
        # Логирование
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален отзыв на продавца: {buyer_info} → {seller_info} (рейтинг: {rating}, дата: {created_date})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Отзыв на продавца успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Admin delete seller review error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)






def admin_logs(request):
    """Страница управления логами действий"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    if current_user.role_id != 1:  # Только администратор
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    # Фильтрация и поиск
    search_query = request.GET.get('search', '')
    user_filter = request.GET.get('user', '')
    action_filter = request.GET.get('action', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    
    # Получаем все логи с предварительной загрузкой данных
    logs = UserActivityLog.objects.all().select_related('user').order_by('-created_at')
    
    # Применяем фильтры
    if search_query:
        logs = logs.filter(
            Q(description__icontains=search_query) |
            Q(ip_address__icontains=search_query) |
            Q(user_agent__icontains=search_query) |
            Q(user__login__icontains=search_query) |
            Q(user__firstname__icontains=search_query) |
            Q(user__surname__icontains=search_query)
        )
    
    if user_filter:
        if user_filter == 'anonymous':
            logs = logs.filter(user__isnull=True)
        else:
            logs = logs.filter(user_id=user_filter)
    
    if action_filter:
        logs = logs.filter(action=action_filter)
    
    if date_from:
        try:
            date_from_obj = datetime.strptime(date_from, '%Y-%m-%d')
            logs = logs.filter(created_at__date__gte=date_from_obj)
        except ValueError:
            pass
    
    if date_to:
        try:
            date_to_obj = datetime.strptime(date_to, '%Y-%m-%d')
            logs = logs.filter(created_at__date__lte=date_to_obj)
        except ValueError:
            pass
    
    # Получаем данные для фильтров
    users_with_logs = Users.objects.filter(useractivitylog__isnull=False).distinct().order_by('login')
    all_users = [('anonymous', 'Анонимные пользователи')] + [(str(u.id_user), u.login) for u in users_with_logs]
    
    actions = UserActivityLog.ACTION_CHOICES
    
    # Статистика
    today = timezone.now().date()
    
    stats = {
        'total_logs': logs.count(),
        'today_logs': logs.filter(created_at__date=today).count(),
        'unique_users': logs.values('user').distinct().count(),
        'top_action': logs.values('action').annotate(
            count=Count('id_log')
        ).order_by('-count').first(),
        'logs_by_hour': logs.filter(
            created_at__date=today
        ).extra({
            'hour': "EXTRACT(HOUR FROM created_at)"
        }).values('hour').annotate(
            count=Count('id_log')
        ).order_by('hour'),
    }
    
    # Пагинация
    paginator = Paginator(logs, 20)  # 20 логов на страницу
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    return render(request, 'admin/logs.html', {
        'logs': page_obj,
        'page_obj': page_obj,
        'all_users': all_users,
        'actions': actions,
        'stats': stats,
        'search_query': search_query,
        'user_filter': user_filter,
        'action_filter': action_filter,
        'date_from': date_from,
        'date_to': date_to,
        'is_admin': current_user.role_id == 1,
        'is_manager': current_user.role_id == 3
    })

@csrf_exempt
@require_POST
def admin_delete_log(request, log_id):
    """Удаление лога (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ может удалять логи
    if current_user.role_id != 1:  # Только администратор
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        log = get_object_or_404(UserActivityLog, id_log=log_id)
        
        action_info = log.get_action_display()
        user_info = log.user.login if log.user else 'Анонимный пользователь'
        created_date = log.created_at.strftime('%d.%m.%Y %H:%M:%S')
        description = log.description[:50] + ('...' if len(log.description) > 50 else '')
        
        # Удаляем лог
        log.delete()
        
        # Логирование удаления лога
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Удален лог #{log_id}: {user_info} - {action_info} - {description} (дата: {created_date})'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Лог успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Admin delete log error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def admin_clear_all_logs(request):
    """Очистка всех логов (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    current_user = Users.objects.get(id_user=user_id)
    
    # Только админ может очищать все логи
    if current_user.role_id != 1:  # Только администратор
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        # Получаем количество логов перед удалением
        logs_count = UserActivityLog.objects.count()
        
        # Удаляем все логи
        deleted_count = UserActivityLog.objects.all().delete()[0]
        
        # Логирование очистки логов
        UserActivityLog.objects.create(
            user=current_user,
            action='admin_action',
            description=f'Очищены все логи. Удалено {deleted_count} записей'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Все логи успешно удалены ({deleted_count} записей)'
        })
        
    except Exception as e:
        logger.error(f'Admin clear all logs error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)



# views.py
from django.db.models import Avg, Count
import json
from django.http import JsonResponse
from django.views.decorators.http import require_POST
from django.shortcuts import get_object_or_404
from .models import SellerReviews, Orders, OrderItems, Users

@require_POST
def add_seller_review(request, seller_id):
    """Добавить отзыв на продавца"""
    user_id = request.session.get('user_id')
    if not user_id:
        return JsonResponse({'error': 'Не авторизован'}, status=401)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный JSON'}, status=400)

    rating = data.get('rating')
    review_text = data.get('review_text')
    order_id = data.get('order_id')

    if not rating or not review_text:
        return JsonResponse({'error': 'Заполните все поля'}, status=400)

    # Проверяем, существует ли продавец
    seller = get_object_or_404(Users, id_user=seller_id)
    
    # Проверяем, что пользователь не оставляет отзыв сам себе
    if user_id == seller_id:
        return JsonResponse({'error': 'Нельзя оставить отзыв самому себе'}, status=400)

    # 🔴 ПРОВЕРКА: вдруг отзыв уже есть
    if SellerReviews.objects.filter(buyer_id=user_id, seller_id=seller_id).exists():
        return JsonResponse({'error': 'Вы уже оставляли отзыв этому продавцу'}, status=400)

    
    # Теперь ищем заказы со статусами 'pending' или 'completed'
    has_purchased = OrderItems.objects.filter(
        order__user_id=user_id,
        order__status__in=['pending', 'completed'],
        product__seller_id=seller_id
    ).exists()

    if not has_purchased:
        # Альтернативная проверка через Orders
        has_purchased = Orders.objects.filter(
            user_id=user_id,
            status__in=['pending', 'completed'],
            orderitems__product__seller_id=seller_id
        ).exists()

    if not has_purchased:
        return JsonResponse({'error': 'Вы можете оставить отзыв только после покупки у этого продавца'}, status=400)

    # Создаем отзыв
    review = SellerReviews.objects.create(
        buyer_id=user_id,
        seller_id=seller_id,
        rating=int(rating),
        review_text=review_text,
        order_id=order_id if order_id else None
    )

    # Обновляем рейтинг продавца
    update_seller_rating(seller_id)

    return JsonResponse({
        'success': True,
        'review_id': review.id_sellerreview
    })


@require_POST
def edit_seller_review(request):
    """Редактировать отзыв на продавца"""
    user_id = request.session.get('user_id')
    if not user_id:
        return JsonResponse({'error': 'Не авторизован'}, status=401)

    try:
        data = json.loads(request.body)
        review_id = data.get('review_id')
        rating = int(data.get('rating'))
        review_text = data.get('review_text')
    except:
        return JsonResponse({'error': 'Неверные данные'}, status=400)

    try:
        review = SellerReviews.objects.get(id_sellerreview=review_id, buyer_id=user_id)
    except SellerReviews.DoesNotExist:
        return JsonResponse({'error': 'Отзыв не найден'}, status=404)

    review.rating = rating
    review.review_text = review_text
    review.save()

    # Обновляем рейтинг продавца
    update_seller_rating(review.seller_id)

    return JsonResponse({'success': True})


@require_POST
def delete_seller_review(request):
    """Удалить отзыв на продавца"""
    user_id = request.session.get('user_id')
    if not user_id:
        return JsonResponse({'error': 'Не авторизован'}, status=401)

    try:
        data = json.loads(request.body)
        review_id = data.get('review_id')
    except:
        return JsonResponse({'error': 'Неверные данные'}, status=400)

    try:
        review = SellerReviews.objects.get(id_sellerreview=review_id, buyer_id=user_id)
        seller_id = review.seller_id
        review.delete()
        
        # Обновляем рейтинг продавца
        update_seller_rating(seller_id)
        
    except SellerReviews.DoesNotExist:
        return JsonResponse({'error': 'Отзыв не найден'}, status=404)

    return JsonResponse({'success': True})


def update_seller_rating(seller_id):
    """Обновить рейтинг продавца"""
    from django.db.models import Avg
    from .models import Users
    
    avg_rating = SellerReviews.objects.filter(
        seller_id=seller_id
    ).aggregate(avg=Avg('rating'))['avg'] or 0
    
    seller = Users.objects.get(id_user=seller_id)
    seller.seller_rating = round(avg_rating, 1)
    seller.save()



def product_detail(request, product_id):
    """Страница товара"""
    product = get_object_or_404(Products, id_product=product_id, is_active=True)
    
    available_tovars = product.tovars.filter(is_sold=False)

    # Увеличиваем просмотры
    in_wishlist = False
    in_cart = False
    cart_quantity = 0
    
    if 'user_id' in request.session:
        user_id = request.session['user_id']
        in_wishlist = Wishlists.objects.filter(user_id=user_id, product=product).exists()
        cart_item = Cart.objects.filter(user_id=user_id, product=product).first()
        if cart_item:
            in_cart = True
            cart_quantity = cart_item.quantity
    
    # Похожие товары
    similar_products = Products.objects.filter(
        category=product.category,
        is_active=True
    ).exclude(id_product=product.id_product).distinct()[:4]
    
    # Отзывы на товар
    product_reviews = ProductReviews.objects.filter(product=product).select_related('user').order_by('-created_at')
    
    # Средний рейтинг товара
    avg_product_rating = product_reviews.aggregate(avg=Avg('rating'))['avg'] or 0
    
    # Отзывы на продавца
    seller_reviews = SellerReviews.objects.filter(seller=product.seller).select_related('buyer').order_by('-created_at')
    
    # Средний рейтинг продавца
    avg_seller_rating = seller_reviews.aggregate(avg=Avg('rating'))['avg'] or 0
    
    # Отзыв текущего пользователя на товар (если есть)
    user_product_review = None
    can_review_product = False
    has_purchased_product = False
    
    # Отзыв текущего пользователя на продавца (если есть)
    user_seller_review = None
    can_review_seller = False
    has_purchased_from_seller = False
    
    if 'user_id' in request.session:
        user_id = request.session['user_id']
        
        # =========== ПРОВЕРКИ ДЛЯ ТОВАРА ===========
        user_product_review = ProductReviews.objects.filter(
            user_id=user_id, 
            product=product
        ).first()
        
        # Проверяем, покупал ли пользователь этот КОНКРЕТНЫЙ товар
        has_purchased_product = OrderItems.objects.filter(
            order__user_id=user_id,
            order__status__in=['pending', 'completed'],
            product=product  # Проверяем именно этот продукт
        ).exists()

        if not has_purchased_product:
            has_purchased_product = Orders.objects.filter(
                user_id=user_id,
                status__in=['pending', 'completed'],
                orderitems__product=product  # Проверяем именно этот продукт
            ).exists()
        
        can_review_product = has_purchased_product and user_product_review is None
        
        # =========== ПРОВЕРКИ ДЛЯ ПРОДАВЦА ===========
        user_seller_review = SellerReviews.objects.filter(
            buyer_id=user_id, 
            seller=product.seller
        ).first()
        
        
        # (а не любой товар этого продавца)
        has_purchased_from_seller = OrderItems.objects.filter(
            order__user_id=user_id,
            order__status__in=['pending', 'completed'],
            product=product,  # Проверяем именно этот продукт
            product__seller=product.seller  # Убеждаемся, что продавец тот же
        ).exists()

        if not has_purchased_from_seller:
            has_purchased_from_seller = Orders.objects.filter(
                user_id=user_id,
                status__in=['pending', 'completed'],
                orderitems__product=product,  # Проверяем именно этот продукт
                orderitems__product__seller=product.seller
            ).exists()
        
        # МОЖНО ОСТАВИТЬ ОТЗЫВ НА ПРОДАВЦА, ЕСЛИ:
        # 1. Купил именно этот товар
        # 2. Еще не оставлял отзыв этому продавцу
        # 3. Не является самим продавцом
        can_review_seller = has_purchased_from_seller and user_seller_review is None and user_id != product.seller.id_user
    
    return render(request, 'products/detail.html', {
        'product': product,
        'available_tovars': available_tovars,
        'available_tovars_count': available_tovars.count(),
        'in_wishlist': in_wishlist,
        'in_cart': in_cart,
        'cart_quantity': cart_quantity,
        'similar_products': similar_products,
        'product_reviews': product_reviews,
        'avg_product_rating': avg_product_rating,
        'seller_reviews': seller_reviews,
        'avg_seller_rating': avg_seller_rating,
        'user_product_review': user_product_review,
        'can_review_product': can_review_product,
        'has_purchased_product': has_purchased_product,
        'user_seller_review': user_seller_review,
        'can_review_seller': can_review_seller,
        'has_purchased_from_seller': has_purchased_from_seller
    })




# ==================== АУТЕНТИФИКАЦИЯ ====================

def password_recovery(request):
    """Восстановление пароля"""
    if request.method == 'POST':
        email = request.POST.get('email')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        errors = []
        
        # Валидация
        if new_password != confirm_password:
            errors.append('Пароли не совпадают')
        
        if len(new_password) < 6:
            errors.append('Пароль должен содержать минимум 6 символов')
        
        try:
            user = Users.objects.get(login=email, is_active=True)
            
            if errors:
                for error in errors:
                    messages.error(request, error)
                return render(request, 'auth/password_recovery.html', {'success': False})
            
            # Меняем пароль
            user.set_password(new_password)
            user.save()
            
            # Завершаем все активные сессии (опционально)
            # Можно добавить логику для завершения сессий
            
            # Логирование
            UserActivityLog.objects.create(
                user=user,
                action='change_password',
                description='Восстановление пароля через email',
                ip_address=get_client_ip(request),
                user_agent=request.META.get('HTTP_USER_AGENT', '')
            )
            
            messages.success(request, 'Пароль успешно изменен!')
            return render(request, 'auth/password_recovery.html', {'success': True})
            
        except Users.DoesNotExist:
            messages.error(request, 'Пользователь с таким email не найден или аккаунт заблокирован')
            return render(request, 'auth/password_recovery.html', {'success': False})
        except Exception as e:
            logger.error(f'Password recovery error: {str(e)}')
            messages.error(request, f'Ошибка при восстановлении пароля: {str(e)}')
            return render(request, 'auth/password_recovery.html', {'success': False})
    
    return render(request, 'auth/password_recovery.html', {'success': False})

def register(request):
    """Регистрация пользователя"""
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        
        # Валидация
        errors = []
        if password != confirm_password:
            errors.append('Пароли не совпадают')
        if len(password) < 6:
            errors.append('Пароль должен содержать минимум 6 символов')
        if not email or '@' not in email:
            errors.append('Введите корректный email')
        
        if Users.objects.filter(login=email).exists():
            errors.append('Пользователь с таким email уже существует')
        
        if errors:
            for error in errors:
                messages.error(request, error)
            return render(request, 'auth/register.html')
        
        try:
            # Создание пользователя
            user = Users.objects.create(
                login=email,
                firstname=first_name,
                surname=last_name,
                role_id=2  # Обычный пользователь
            )
            user.set_password(password)
            user.save()
            
            # Автоматический вход
            request.session['user_id'] = user.id_user
            request.session['user_login'] = user.login
            request.session['user_name'] = f"{user.firstname} {user.surname}"
            request.session['user_role'] = user.role.role_name
            request.session['user_role_id'] = user.role_id
            
            # Логирование
            UserActivityLog.objects.create(
                user=user,
                action='register',
                description=f'Регистрация нового пользователя {email}'
            )
            
            messages.success(request, 'Регистрация успешна! Добро пожаловать!')
            return redirect('home')
            
        except Exception as e:
            messages.error(request, f'Ошибка при регистрации: {str(e)}')
            logger.error(f'Registration error: {str(e)}')
    
    return render(request, 'auth/register.html')

def login_view(request):
    """Вход в систему"""
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')
        
        try:
            user = Users.objects.get(login=email, is_active=True)
            
            if user.check_password(password):
                # Сохраняем в сессии
                request.session['user_id'] = user.id_user
                request.session['user_login'] = user.login
                request.session['user_name'] = f"{user.firstname} {user.surname}"
                request.session['user_role'] = user.role.role_name
                request.session['user_role_id'] = user.role_id
                
                # Обновляем last_login
                user.last_login = timezone.now()
                user.save()
                
                # Логирование
                UserActivityLog.objects.create(
                    user=user,
                    action='login',
                    description=f'Вход в систему с IP: {get_client_ip(request)}',
                    ip_address=get_client_ip(request),
                    user_agent=request.META.get('HTTP_USER_AGENT', '')
                )
                
                messages.success(request, f'Добро пожаловать, {user.firstname}!')
                
                # Редирект в зависимости от роли
                next_url = request.GET.get('next', 'home')
                if user.role_id in [1, 3]:  # Админ или менеджер
                    next_url = 'admin_dashboard'
                
                return redirect(next_url)
            else:
                messages.error(request, 'Неверный пароль')
                
        except Users.DoesNotExist:
            messages.error(request, 'Пользователь не найден или аккаунт заблокирован')
        except Exception as e:
            messages.error(request, f'Ошибка входа: {str(e)}')
            logger.error(f'Login error: {str(e)}')
    
    return render(request, 'auth/login.html')

def logout_view(request):
    """Выход из системы"""
    user_id = request.session.get('user_id')
    if user_id:
        try:
            user = Users.objects.get(id_user=user_id)
            # Логирование
            UserActivityLog.objects.create(
                user=user,
                action='logout',
                description='Выход из системы',
                ip_address=get_client_ip(request)
            )
        except Users.DoesNotExist:
            pass
    
    # Очищаем сессию
    request.session.flush()
    
    messages.success(request, 'Вы успешно вышли из системы')
    return redirect('home')

# ==================== ПРОФИЛЬ ПОЛЬЗОВАТЕЛЯ ====================

def profile(request):
    """Профиль пользователя"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Статистика
    orders_count = Orders.objects.filter(user=user).count()
    wishlist_count = Wishlists.objects.filter(user=user).count()
    reviews_count = ProductReviews.objects.filter(user=user).count()
    
    # Последние заказы
    recent_orders = Orders.objects.filter(user=user).order_by('-order_created_at')[:5]
    
    # Активность
    recent_activity = UserActivityLog.objects.filter(user=user).order_by('-created_at')[:10]
    
    

    return render(request, 'user/profile.html', {
        'user_obj': user,
        'orders_count': orders_count,
        'wishlist_count': wishlist_count,
        'reviews_count': reviews_count,
        'recent_orders': recent_orders, 
        'recent_activity': recent_activity
    })

@require_POST
def update_profile(request):
    """Обновление профиля"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        firstname = request.POST.get('firstname')
        surname = request.POST.get('surname')
        
        if firstname and surname:
            old_name = f"{user.firstname} {user.surname}"
            user.firstname = firstname
            user.surname = surname
            user.save()
            
            # Обновляем сессию
            request.session['user_name'] = f"{firstname} {surname}"
            
            # Логирование
            UserActivityLog.objects.create(
                user=user,
                action='update_profile',
                description=f'Обновление профиля: {old_name} -> {firstname} {surname}'
            )
            
            messages.success(request, 'Профиль обновлен')
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'error': 'Все поля обязательны'}, status=400)
            
    except Exception as e:
        logger.error(f'Update profile error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@require_POST
def change_password(request):
    """Смена пароля"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        current_password = request.POST.get('current_password')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        if not user.check_password(current_password):
            return JsonResponse({'error': 'Текущий пароль неверен'}, status=400)
        
        if new_password != confirm_password:
            return JsonResponse({'error': 'Пароли не совпадают'}, status=400)
        
        if len(new_password) < 6:
            return JsonResponse({'error': 'Пароль должен содержать минимум 6 символов'}, status=400)
        
        user.set_password(new_password)
        user.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='change_password',
            description='Смена пароля'
        )
        
        messages.success(request, 'Пароль успешно изменен')
        return JsonResponse({'success': True})
            
    except Exception as e:
        logger.error(f'Change password error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)









# ==================== КОРЗИНА ====================


def cart_view(request):
    """Корзина пользователя"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для просмотра корзины необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    
    try:
        user = Users.objects.get(id_user=user_id)
    except Users.DoesNotExist:
        messages.error(request, 'Пользователь не найден')
        return redirect('login')
    
    cart_items = Cart.objects.filter(user_id=user_id).select_related('product')
    
    total_price = 0
    cart_items_with_total = []
    
    for item in cart_items:
        item_total = item.product.price * item.quantity
        total_price += item_total
        
        cart_items_with_total.append({
            'item': item,
            'item_total': item_total
        })
    
    # Проверяем, есть ли сохраненный промокод в сессии
    applied_promocode_data = request.session.get('applied_promocode')
    final_total = total_price
    discount_percent = 0
    discount_amount = 0
    
    if applied_promocode_data:
        try:
            # Получаем промокод из БД
            promocode = PromoCodes.objects.get(code=applied_promocode_data['code'])
            if promocode.is_valid:
                # Применяем скидку
                final_total = promocode.apply_discount(total_price)
                discount_amount = total_price - final_total
                discount_percent = promocode.discount_percent
        except PromoCodes.DoesNotExist:
            # Если промокод не найден, удаляем из сессии
            request.session.pop('applied_promocode', None)
    
    return render(request, 'cart/index.html', {
        'cart_items': cart_items,
        'cart_items_with_total': cart_items_with_total,
        'total_price': total_price,
        'final_price': final_total,  # Добавляем финальную цену со скидкой
        'discount_amount': discount_amount,
        'discount_percent': discount_percent,
        'user_balance': user.balance,
        'user_obj': user
    })

@csrf_exempt
@require_POST
def add_to_cart(request):
    """Добавление в корзину (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        data = json.loads(request.body)
        product_id = data.get('product_id')
        quantity = int(data.get('quantity', 1))
        
        product = get_object_or_404(Products, id_product=product_id, is_active=True)
        
        # Проверка наличия
        available_count = product.available_tovars_count
        if available_count < quantity:
            return JsonResponse({
                'error': f'Недостаточно товаров. Доступно: {available_count} шт.'
            }, status=400)
        
        # Проверяем, есть ли уже товар в корзине
        cart_item = Cart.objects.filter(user_id=user_id, product=product).first()
        
        available_count = product.available_tovars_count  # Используем property
        if available_count < quantity:
            return JsonResponse({
                'error': f'Недостаточно товаров. Доступно: {available_count} шт.'
            }, status=400)

        if cart_item:
            # Если товар уже есть, обновляем количество
            new_quantity = cart_item.quantity + quantity
            if product.stock_quantity < new_quantity:
                return JsonResponse({'error': 'Недостаточно товара на складе'}, status=400)
            
            cart_item.quantity = new_quantity
            cart_item.save()
            action = 'updated'
        else:
            # Если товара нет, создаем новую запись
            Cart.objects.create(user_id=user_id, product=product, quantity=quantity)
            action = 'added'
        
        # Логирование
        UserActivityLog.objects.create(
            user_id=user_id,
            action='add_to_cart',
            description=f'Добавление в корзину: {product.title} (x{quantity})'
        )
        
        cart_count = Cart.objects.filter(user_id=user_id).count()
        
        return JsonResponse({
            'success': True,
            'message': 'Товар добавлен в корзину',
            'cart_count': cart_count,
            'action': action,
            'item': {
                'id': cart_item.id_cart if cart_item else '',
                'product_id': product.id_product,
                'title': product.title,
                'price': float(product.price),
                'quantity': cart_item.quantity if cart_item else quantity,
                'image': product.image_url
            }
        })
        
    except Exception as e:
        logger.error(f'Add to cart error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def update_cart_item(request, item_id):
    """Обновление количества товара в корзине"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        cart_item = Cart.objects.get(id_cart=item_id, user_id=user_id)
        data = json.loads(request.body)
        quantity = int(data.get('quantity', 1))
        
        if quantity <= 0:
            cart_item.delete()
            action = 'removed'
        else:
            if cart_item.product.available_tovars_count < quantity:
                return JsonResponse({'error': 'Недостаточно товара на складе'}, status=400)
            
            cart_item.quantity = quantity
            cart_item.save()
            action = 'updated'
        
        # Пересчет общей суммы
        cart_items = Cart.objects.filter(user_id=user_id).select_related('product')
        total_price = sum(item.product.price * item.quantity for item in cart_items)
        
        return JsonResponse({
            'success': True,
            'action': action,
            'total_price': float(total_price),
            'item_price': float(cart_item.product.price * cart_item.quantity) if action == 'updated' else 0
        })
        
    except Cart.DoesNotExist:
        return JsonResponse({'error': 'Товар не найден в корзине'}, status=404)
    except Exception as e:
        logger.error(f'Update cart error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def remove_from_cart(request, item_id):
    """Удаление из корзины"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        cart_item = Cart.objects.get(id_cart=item_id, user_id=user_id)
        product_title = cart_item.product.title
        cart_item.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user_id=user_id,
            action='remove_from_cart',
            description=f'Удаление из корзины: {product_title}'
        )
        
        # Пересчет
        cart_items = Cart.objects.filter(user_id=user_id).select_related('product')
        total_price = sum(item.product.price * item.quantity for item in cart_items)
        cart_count = cart_items.count()
        
        return JsonResponse({
            'success': True,
            'message': 'Товар удален из корзины',
            'total_price': float(total_price),
            'cart_count': cart_count
        })
        
    except Cart.DoesNotExist:
        return JsonResponse({'error': 'Товар не найден в корзине'}, status=404)
    except Exception as e:
        logger.error(f'Remove from cart error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


# appip/views.py

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def reserve_tovar(request):
    """Резервирование товара перед покупкой"""
    product_id = request.data.get('product_id')
    tovar_id = request.data.get('tovar_id')
    
    try:
        tovar = Tovars.objects.get(id_tovar=tovar_id, is_sold=False)
        product = Products.objects.get(id_product=product_id)
        
        # Проверяем, что товар относится к продукту
        if not product.tovars.filter(id_tovar=tovar_id).exists():
            return Response({'error': 'Товар не принадлежит продукту'}, 
                          status=status.HTTP_400_BAD_REQUEST)
        
        # Резервируем товар
        tovar.is_sold = True
        tovar.sold_at = timezone.now()
        tovar.save()
        
        return Response({
            'success': True,
            'message': 'Товар зарезервирован',
            'tovar_id': tovar.id_tovar
        })
        
    except Tovars.DoesNotExist:
        return Response({'error': 'Товар не найден или уже продан'}, 
                      status=status.HTTP_404_NOT_FOUND)

@api_view(['GET'])
def product_tovars(request, product_id):
    """Получение товаров для продукта"""
    product = get_object_or_404(Products, id_product=product_id)
    tovars = product.tovars.filter(is_sold=False).order_by('-created_at')
    
    serializer = TovarsSerializer(tovars, many=True)
    return Response({
        'product_id': product_id,
        'title': product.title,
        'available_tovars': serializer.data,
        'count': tovars.count()
    })

# ==================== ВИШЛИСТ ====================

def wishlist_view(request):
    """Вишлист пользователя"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для просмотра избранного необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    wishlist_items = Wishlists.objects.filter(user_id=user_id).select_related('product')
    
    total_price = sum(item.product.price for item in wishlist_items)
    avg_rating = wishlist_items.aggregate(avg=Avg('product__rating'))['avg'] or 0
    
    return render(request, 'wishlist/index.html', {
        'wishlist_items': wishlist_items,
        'total_price': total_price,
        'avg_rating': avg_rating
    })

@csrf_exempt
@require_POST
def toggle_wishlist(request):
    """Добавление/удаление из вишлиста (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        data = json.loads(request.body)
        product_id = data.get('product_id')
        
        product = get_object_or_404(Products, id_product=product_id, is_active=True)
        wishlist_item = Wishlists.objects.filter(user_id=user_id, product=product).first()
        
        if wishlist_item:
            wishlist_item.delete()
            action = 'removed'
            message = 'Товар удален из избранного'
        else:
            Wishlists.objects.create(user_id=user_id, product=product)
            action = 'added'
            message = 'Товар добавлен в избранное'
        
        # Логирование
        UserActivityLog.objects.create(
            user_id=user_id,
            action='add_to_wishlist' if action == 'added' else 'remove_from_wishlist',
            description=f'{message}: {product.title}'
        )
        
        wishlist_count = Wishlists.objects.filter(user_id=user_id).count()
        
        return JsonResponse({
            'success': True,
            'action': action,
            'message': message,
            'wishlist_count': wishlist_count
        })
        
    except Exception as e:
        logger.error(f'Toggle wishlist error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)







# ==================== ЗАКАЗЫ ====================

def orders_view(request):
    """История заказов"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для просмотра заказов необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    orders = Orders.objects.filter(user_id=user_id).select_related('user').prefetch_related(
        'orderitems_set__product'
    ).order_by('-order_created_at')
    
    return render(request, 'orders/index.html', {
        'orders': orders
    })

def order_detail(request, order_id):
    """Детали заказа"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для просмотра заказа необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    order = get_object_or_404(Orders, id_order=order_id, user_id=user_id)
    
    # Загружаем товары через tovar
    order_items = OrderItems.objects.filter(order=order).select_related('product', 'tovar')
    
    # Проверяем, можно ли оспорить заказ
    can_dispute = False
    dispute_deadline = None
    
    if order.status == 'completed':
        # Проверяем, есть ли уже чат
        existing_chat = order.chats_set.first()
        if not existing_chat:
            # Вычисляем время подтверждения заказа
            from .models import UserActivityLog
            from django.utils import timezone
            from django.conf import settings
            
            confirm_log = UserActivityLog.objects.filter(
                user_id=user_id,
                action='create_order',
                description__icontains=f'Заказ №{order.id_order} подтвержден',
                created_at__gte=order.order_created_at
            ).order_by('-created_at').first()
            
            if confirm_log:
                confirm_time = confirm_log.created_at
            else:
                confirm_time = order.order_created_at
            
            # Приводим confirm_time к timezone-aware
            if timezone.is_naive(confirm_time):
                confirm_time = timezone.make_aware(confirm_time, timezone=timezone.get_current_timezone())
            
            now = timezone.now()
            time_since_confirm = now - confirm_time
            can_dispute = time_since_confirm.total_seconds() < 24 * 3600
            dispute_deadline = confirm_time + timedelta(hours=24)
    
    return render(request, 'orders/detail.html', {
        'order': order,
        'order_items': order_items,
        'can_dispute': can_dispute,
        'dispute_deadline': dispute_deadline
    })


@csrf_exempt
@require_POST
def create_order(request):
    """Создание заказа"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        
        cart_items = Cart.objects.filter(user_id=user_id).select_related('product')
        
        if not cart_items:
            return JsonResponse({'error': 'Корзина пуста'}, status=400)
        
        # Считаем общую сумму
        total_cost = Decimal('0')
        for cart_item in cart_items:
            total_cost += cart_item.product.price * Decimal(str(cart_item.quantity))
        
        # Проверяем наличие промокода в сессии
        applied_promocode_data = request.session.get('applied_promocode')
        final_cost = total_cost
        
        if applied_promocode_data:
            try:
                promocode = PromoCodes.objects.get(code=applied_promocode_data['code'])
                if promocode.is_valid:
                    final_cost = promocode.apply_discount(total_cost)
                    # Увеличиваем счетчик использований
                    promocode.used_count += 1
                    promocode.save()
                    # Удаляем промокод из сессии
                    request.session.pop('applied_promocode', None)
            except PromoCodes.DoesNotExist:
                request.session.pop('applied_promocode', None)
        
        # Проверяем наличие товаров
        for cart_item in cart_items:
            available_count = cart_item.product.available_tovars_count
            if available_count < cart_item.quantity:
                return JsonResponse({
                    'error': f'Недостаточно товара "{cart_item.product.title}" на складе. Доступно: {available_count} шт.'
                }, status=400)
        
        # Проверяем баланс (используем final_cost со скидкой)
        if user.balance < final_cost:
            return JsonResponse({'error': 'Недостаточно средств на балансе'}, status=400)
        
        # Собираем информацию о продавцах для логирования
        sellers_info = {}
        for cart_item in cart_items:
            item_total = cart_item.product.price * cart_item.quantity
            seller_id = cart_item.product.seller_id
            if seller_id not in sellers_info:
                sellers_info[seller_id] = {
                    'login': cart_item.product.seller.login,
                    'total': 0
                }
            sellers_info[seller_id]['total'] += item_total
        
        # Создаем заказ с итоговой суммой со скидкой
        order = Orders.objects.create(
            user=user,
            total_cost=final_cost,  # Используем сумму со скидкой!
            status='pending',
            payment_method='balance',
            payment_reference=f'ORDER_{user_id}_{timezone.now().strftime("%Y%m%d%H%M%S")}'
        )
        
        # Создаем элементы заказа
        for cart_item in cart_items:
            available_tovars = cart_item.product.tovars.filter(is_sold=False)[:cart_item.quantity]
            
            if len(available_tovars) < cart_item.quantity:
                return JsonResponse({
                    'error': f'Недостаточно доступных товаров для "{cart_item.product.title}"'
                }, status=400)
            
            for tovar in available_tovars:
                # Вычисляем цену товара с учетом скидки (пропорционально)
                # Общая скидка распределяется пропорционально стоимости товаров
                if applied_promocode_data and final_cost < total_cost:
                    # Доля этого товара в общей сумме
                    item_original_price = cart_item.product.price
                    item_discounted_price = item_original_price * (final_cost / total_cost)
                    price_to_use = item_discounted_price
                else:
                    price_to_use = cart_item.product.price
                
                order_item = OrderItems.objects.create(
                    order=order,
                    product=cart_item.product,
                    tovar=tovar,
                    quantity=1,
                    price_at_time_of_purchase=price_to_use,  # Сохраняем цену со скидкой
                    status='pending'
                )
                
                tovar.is_sold = True
                tovar.sold_at = timezone.now()
                tovar.save()
        
        # Списание средств с покупателя (используем сумму со скидкой)
        user.balance -= final_cost
        user.save()
        
        # Создаем транзакцию для покупателя
        Transactions.objects.create(
            user=user,
            order=order,
            amount=final_cost,
            transaction_type='purchase',
            status='completed',
            reference=f'ORDER_{order.id_order}'
        )
        
        # Очищаем корзину
        cart_items.delete()
        
        # Логирование
        discount_info = f" (со скидкой {final_cost} ₽ из {total_cost} ₽)" if applied_promocode_data else ""
        UserActivityLog.objects.create(
            user=user,
            action='create_order',
            description=f'Создан заказ №{order.id_order} на сумму {final_cost} руб.{discount_info}',
            ip_address=get_client_ip(request)
        )
        
        # Логирование для продавцов
        for seller_id, info in sellers_info.items():
            seller = Users.objects.get(id_user=seller_id)
            UserActivityLog.objects.create(
                user=seller,
                action='create_order',
                description=f'Новый заказ №{order.id_order} от {user.login} на сумму {info["total"]} руб. (с учетом скидки)'
            )
        
        return JsonResponse({
            'success': True,
            'message': 'Заказ успешно создан!',
            'order_id': order.id_order,
            'order_total': float(final_cost),
            'original_total': float(total_cost),
            'discount_applied': applied_promocode_data is not None
        })
        
    except Exception as e:
        logger.error(f'Create order error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


@require_GET
def get_unread_chats_count(request):
    """API для получения количества непрочитанных сообщений"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        # Получаем все чаты пользователя
        chats_as_buyer = Chats.objects.filter(buyer_id=user_id, is_active=True)
        chats_as_seller = Chats.objects.filter(seller_id=user_id, is_active=True)
        
        total_unread = 0
        buyer_unread = 0
        seller_unread = 0
        chats_data = []
        
        # Считаем непрочитанные сообщения в чатах как покупателя
        for chat in chats_as_buyer:
            unread_count = Messages.objects.filter(
                chat=chat,
                sender_id=chat.seller_id,
                is_read=False
            ).count()
            
            if unread_count > 0:
                total_unread += unread_count
                buyer_unread += unread_count
                chats_data.append({
                    'chat_id': chat.id_chat,
                    'role': 'buyer',
                    'unread_count': unread_count
                })
        
        # Считаем непрочитанные сообщения в чатах как продавца
        for chat in chats_as_seller:
            unread_count = Messages.objects.filter(
                chat=chat,
                sender_id=chat.buyer_id,
                is_read=False
            ).count()
            
            if unread_count > 0:
                total_unread += unread_count
                seller_unread += unread_count
                chats_data.append({
                    'chat_id': chat.id_chat,
                    'role': 'seller',
                    'unread_count': unread_count
                })
        
        return JsonResponse({
            'success': True,
            'total_unread': total_unread,
            'chats': {
                'buyer_unread': buyer_unread,
                'seller_unread': seller_unread,
                'list': chats_data
            }
        })
        
    except Exception as e:
        logger.error(f'Get unread chats count error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)





@csrf_exempt
@require_POST
def confirm_order(request, order_id):
    """Подтверждение получения заказа (покупателем)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        order = get_object_or_404(Orders, id_order=order_id, user_id=user_id)
        
        # Проверяем текущий статус
        if order.status != 'pending':
            return JsonResponse({
                'error': f'Заказ уже имеет статус "{order.get_status_display()}"'
            }, status=400)
        
        # Получаем все товары в заказе
        order_items = OrderItems.objects.filter(order=order).select_related('product__seller')
        
        # Группируем по продавцам для начисления средств
        seller_payments = {}
        for item in order_items:
            seller = item.product.seller
            # Сумма за этот товар (уже со скидкой)
            amount = item.price_at_time_of_purchase * item.quantity
            
            if seller.id_user not in seller_payments:
                seller_payments[seller.id_user] = {
                    'seller': seller,
                    'amount': 0
                }
            seller_payments[seller.id_user]['amount'] += amount
        
        # Начисляем деньги каждому продавцу
        for seller_data in seller_payments.values():
            seller = seller_data['seller']
            amount = seller_data['amount']
            
            # Начисляем средства
            seller.balance += amount
            seller.save()
            
            # Создаем транзакцию для продавца
            Transactions.objects.create(
                user=seller,
                order=order,
                amount=amount,
                transaction_type='sale',
                status='completed',
                reference=f'SALE_ORDER_{order.id_order}_SELLER_{seller.id_user}'
            )
            
            # Логирование для продавца
            UserActivityLog.objects.create(
                user=seller,
                action='create_order',
                description=f'Зачисление средств за заказ №{order.id_order}: +{amount} ₽'
            )
        
        # Меняем статус на "Завершен"
        old_status = order.status
        order.status = 'completed'
        order.save()
        
        # Обновляем статусы всех элементов заказа
        order_items.update(status='delivered')
        
        # Логирование для покупателя
        UserActivityLog.objects.create(
            user=user,
            action='create_order',
            description=f'Заказ №{order.id_order} подтвержден пользователем (статус изменен с {old_status} на completed)',
            ip_address=get_client_ip(request)
        )
        
        # Формируем информацию о выплатах для ответа
        payments_info = []
        for seller_data in seller_payments.values():
            payments_info.append({
                'seller_name': seller_data['seller'].login,
                'amount': float(seller_data['amount'])
            })
        
        return JsonResponse({
            'success': True,
            'message': 'Заказ успешно подтвержден! Средства зачислены продавцам.',
            'new_status': 'completed',
            'status_display': 'Завершен',
            'payments': payments_info
        })
        
    except Exception as e:
        logger.error(f'Confirm order error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


import uuid
from yookassa import Configuration, Payment
from django.conf import settings
from django.urls import reverse

# Настройка ЮKassa
Configuration.account_id = settings.YOOKASSA_SHOP_ID
Configuration.secret_key = settings.YOOKASSA_SECRET_KEY

@csrf_exempt
@require_POST
def deposit_balance(request):
    """Пополнение баланса через ЮMoney"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        data = json.loads(request.body)
        amount = float(data.get('amount', 0))
        
        # Валидация суммы
        if amount < 10:
            return JsonResponse({'error': 'Минимальная сумма пополнения 10 ₽'}, status=400)
        
        if amount > 100000:
            return JsonResponse({'error': 'Максимальная сумма пополнения 100 000 ₽'}, status=400)
        
        # Создаем транзакцию со статусом pending
        transaction = Transactions.objects.create(
            user=user,
            amount=amount,
            transaction_type='deposit',
            status='pending',
            reference=f'DEPOSIT_{user_id}_{int(timezone.now().timestamp())}'
        )
        
        # Создаем платеж в ЮKassa
        idempotence_key = str(uuid.uuid4())
        
        # ВАЖНО: Укажите правильный URL для уведомлений
        notification_url = request.build_absolute_uri(reverse('payment_notification'))
        
        payment = Payment.create({
            "amount": {
                "value": f"{amount:.2f}",
                "currency": "RUB"
            },
            "confirmation": {
                "type": "redirect",
                "return_url": request.build_absolute_uri(reverse('payment_success')) + f'?transaction_id={transaction.id_transaction}'
            },
            "capture": True,
            "description": f"Пополнение баланса пользователя {user.login}",
            "metadata": {
                "user_id": user_id,
                "transaction_id": transaction.id_transaction
            }
        }, idempotence_key)
        
        # Обновляем reference транзакции на ID платежа в ЮKassa
        transaction.reference = payment.id
        transaction.save()
        
        return JsonResponse({
            'success': True,
            'payment_url': payment.confirmation.confirmation_url,
            'payment_id': payment.id,
            'transaction_id': transaction.id_transaction
        })
        
    except Exception as e:
        logger.error(f'Deposit error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_POST
def payment_notification(request):
    """Обработка уведомлений от ЮKassa (webhook)"""
    try:
        # Получаем данные от ЮKassa
        data = json.loads(request.body)
        
        # Логируем для отладки
        logger.info(f"Payment notification received: {data}")
        
        # Проверяем, что это уведомление о платеже
        if data.get('event') == 'payment.succeeded':
            payment_id = data['object']['id']
            
            # Находим транзакцию
            transaction = Transactions.objects.filter(reference=payment_id).first()
            
            if transaction and transaction.status == 'pending':
                # Обновляем статус транзакции
                transaction.status = 'completed'
                transaction.save()
                
                # Начисляем средства на баланс пользователя
                user = transaction.user
                user.balance += transaction.amount
                user.save()
                
                # Логирование
                UserActivityLog.objects.create(
                    user=user,
                    action='deposit',
                    description=f'Пополнение баланса на {transaction.amount} ₽ через ЮMoney',
                    ip_address=get_client_ip(request)
                )
                
                logger.info(f"Balance updated for user {user.id_user}: +{transaction.amount}")
                
                return JsonResponse({'success': True})
            else:
                logger.warning(f"Transaction not found or already completed: {payment_id}")
        
        return JsonResponse({'success': True})
        
    except Exception as e:
        logger.error(f'Payment notification error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_GET
def payment_success(request):
    """Страница успешной оплаты"""
    transaction_id = request.GET.get('transaction_id')
    
    if transaction_id:
        try:
            transaction = Transactions.objects.get(id_transaction=transaction_id)
            
            # Проверяем статус транзакции
            if transaction.status == 'completed':
                messages.success(request, f'Баланс успешно пополнен на {transaction.amount} ₽!')
            elif transaction.status == 'pending':
                # Если транзакция еще в обработке, проверяем статус в ЮKassa
                try:
                    payment = Payment.find_one(transaction.reference)
                    if payment.status == 'succeeded':
                        transaction.status = 'completed'
                        transaction.save()
                        
                        user = transaction.user
                        user.balance += transaction.amount
                        user.save()
                        
                        messages.success(request, f'Баланс успешно пополнен на {transaction.amount} ₽!')
                    else:
                        messages.info(request, 'Платеж обрабатывается. Средства поступят в течение нескольких минут.')
                except:
                    messages.info(request, 'Платеж обрабатывается. Средства поступят в течение нескольких минут.')
        except Transactions.DoesNotExist:
            messages.error(request, 'Транзакция не найдена')
    
    return redirect('profile')


@csrf_exempt
@require_GET
def check_payment_status(request, transaction_id):
    """Проверка статуса платежа (для отладки)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    try:
        transaction = get_object_or_404(Transactions, id_transaction=transaction_id)
        
        # Проверяем, что транзакция принадлежит пользователю
        if transaction.user_id != request.session['user_id']:
            return JsonResponse({'error': 'Доступ запрещен'}, status=403)
        
        # Если транзакция еще в обработке, проверяем в ЮKassa
        if transaction.status == 'pending' and transaction.reference:
            try:
                payment = Payment.find_one(transaction.reference)
                if payment.status == 'succeeded':
                    transaction.status = 'completed'
                    transaction.save()
                    
                    user = transaction.user
                    user.balance += transaction.amount
                    user.save()
            except:
                pass
        
        return JsonResponse({
            'success': True,
            'status': transaction.status,
            'amount': float(transaction.amount),
            'balance': float(transaction.user.balance)
        })
        
    except Exception as e:
        logger.error(f'Check payment error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_GET
def payment_fail(request):
    """Страница неудачной оплаты"""
    messages.error(request, 'Оплата не прошла. Попробуйте снова.')
    return redirect('profile')







import requests
import random
from django.utils import timezone
from .models import TelegramManager, ChatSync

def notify_vk_managers(chat, message, sender):
    """Отправить уведомление менеджерам в VK"""
    try:
        vk_token = settings.VK_GROUP_TOKEN
        group_id = settings.VK_GROUP_ID
        
        if not vk_token:
            logger.error("VK token not provided")
            return
        
        # Получаем всех активных менеджеров
        managers = TelegramManager.objects.filter(is_active=True)
        
        for manager in managers:
            # Проверяем, что это чат этого менеджера
            if chat.seller_id != manager.manager_id:
                continue
            
            # Проверяем наличие vk_peer_id
            if not manager.vk_peer_id:
                logger.info(f"Manager {manager.manager.login} has no VK peer_id")
                continue
            
            # Текст сообщения для VK
            text = f"""💬 НОВОЕ СООБЩЕНИЕ ОТ ПОЛЬЗОВАТЕЛЯ

👤 От: {sender.login}
📦 Товар: {chat.product.title if chat.product else 'Общий чат'}
🆔 Чат ID: {chat.id_chat}

📝 Сообщение:
{message.message_text}

----------------------
Ответьте на это сообщение, чтобы отправить ответ пользователю"""
            
            # Отправляем через VK API
            url = 'https://api.vk.com/method/messages.send'
            random_id = random.randint(-2**31, 2**31 - 1)
            
            params = {
                'peer_id': manager.vk_peer_id,
                'message': text,
                'random_id': random_id,
                'access_token': vk_token,
                'v': '5.131'
            }
            
            response = requests.post(url, data=params, timeout=10)
            
            if response.ok:
                result = response.json()
                if 'error' in result:
                    logger.error(f"VK API Error: {result['error']}")
                else:
                    logger.info(f"Message sent to VK manager {manager.manager.login}")
            else:
                logger.error(f"VK send error: {response.text}")
                
    except Exception as e:
        logger.error(f"VK notification error: {str(e)}")



# ==================== УПРАВЛЕНИЕ ТОВАРАМИ ПОЛЬЗОВАТЕЛЯ ====================

def user_products(request):
    """Страница управления товарами пользователя"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для управления товарами необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Получаем все товары пользователя
    user_products = Products.objects.filter(seller=user).order_by('-created_at')
    
    # Получаем количество отзывов на продавца
    from django.db.models import Count, Avg
    from django.db.models.functions import Coalesce
    
    seller_reviews_count = SellerReviews.objects.filter(seller=user).count()
    
    # Получаем средний рейтинг продавца через агрегацию (как в products_list)
    seller_avg = SellerReviews.objects.filter(seller=user).aggregate(
        avg=Coalesce(Avg('rating'), 0.0)
    )['avg']
    
    # Для отладки - выведем в консоль
    print(f"User: {user.login}, Seller rating from field: {user.seller_rating}, Calculated avg: {seller_avg}, Reviews count: {seller_reviews_count}")
    
    return render(request, 'user/products.html', {
        'user_obj': user,
        'products': user_products,
        'seller_rating': seller_avg,  # Используем вычисленное значение
        'seller_reviews_count': seller_reviews_count,
        'categories': Categories.objects.all(),
        'product_types': ProductTypes.objects.all()
    })

def create_product(request):
    """Создание нового товара"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для создания товара необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    if request.method == 'POST':
        try:
            # Получаем данные из формы
            title = request.POST.get('title')
            description = request.POST.get('description')
            price = request.POST.get('price')
            category_id = request.POST.get('category')
            product_type_id = request.POST.get('product_type')
            
            # ===== НОВЫЙ КОД: Обработка загруженного изображения =====
            product_image = request.FILES.get('product_image')
            main_image_url = None
            
            if product_image:
                # Генерируем имя файла
                import os
                from django.utils.text import slugify
                
                # Создаем безопасное имя файла
                file_extension = os.path.splitext(product_image.name)[1]
                filename = f"{slugify(title)}_{user_id}_{timezone.now().timestamp()}{file_extension}"
                
                # Сохраняем файл
                from django.core.files.storage import default_storage
                from django.core.files.base import ContentFile
                
                file_path = default_storage.save(f'products/{filename}', ContentFile(product_image.read()))
                main_image_url = filename
            # ===== КОНЕЦ НОВОГО КОДА =====
            
            # Получаем товары из скрытого поля
            tovars_data = request.POST.get('tovars_data', '[]')
            tovar_texts = json.loads(tovars_data)
            
            # Валидация
            errors = []
            if not title or len(title.strip()) < 5:
                errors.append('Название товара должно содержать минимум 5 символов')
            if not description or len(description.strip()) < 10:
                errors.append('Описание должно содержать минимум 10 символов')
            if not price or float(price) <= 0:
                errors.append('Цена должна быть положительным числом')
            if not tovar_texts:
                errors.append('Добавьте хотя бы один товар для продажи')
            
            if errors:
                for error in errors:
                    messages.error(request, error)
                return render(request, 'user/create_product.html', {
                    'user_obj': user,
                    'categories': Categories.objects.all(),
                    'product_types': ProductTypes.objects.all(),
                    'form_data': request.POST
                })
            
            # Создаем продукт с изображением
            product = Products.objects.create(
                title=title.strip(),
                description=description.strip(),
                price=price,
                seller=user,
                category_id=category_id if category_id else None,
                product_type_id=product_type_id,
                main_image_url=main_image_url,  # Добавляем имя файла
                is_active=True,
                rating=0
            )
            
            # Добавляем товары
            for text in tovar_texts:
                if text.strip():
                    tovar = Tovars.objects.create(
                        tovar_text=text.strip(),
                        is_sold=False
                    )
                    ProductsTovars.objects.create(
                        product=product,
                        tovar=tovar
                    )
            
            # Логирование
            UserActivityLog.objects.create(
                user=user,
                action='add_product',
                description=f'Создан новый товар: {product.title} с {len(tovar_texts)} товарами'
            )
            
            messages.success(request, f'Товар "{product.title}" успешно создан с {len(tovar_texts)} товарами!')
            return redirect('user_products')
            
        except Exception as e:
            logger.error(f'Create product error: {str(e)}')
            messages.error(request, f'Ошибка при создании товара: {str(e)}')
            return render(request, 'user/create_product.html', {
                'user_obj': user,
                'categories': Categories.objects.all(),
                'product_types': ProductTypes.objects.all(),
                'form_data': request.POST
            })
    
    return render(request, 'user/create_product.html', {
        'user_obj': user,
        'categories': Categories.objects.all(),
        'product_types': ProductTypes.objects.all()
    })

def edit_product(request, product_id):
    """Редактирование товара"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для редактирования товара необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Проверяем, что товар принадлежит пользователю
    product = get_object_or_404(Products, id_product=product_id, seller=user)
    
    if request.method == 'POST':
        try:
            # Получаем данные из формы
            title = request.POST.get('title')
            description = request.POST.get('description')
            price = request.POST.get('price')
            category_id = request.POST.get('category')
            product_type_id = request.POST.get('product_type')
           
            
            # ===== НОВЫЙ КОД: Обработка загруженного изображения =====
            product_image = request.FILES.get('product_image')
            
            if product_image:
                # Генерируем имя файла
                import os
                from django.utils.text import slugify
                
                # Удаляем старое изображение, если оно есть
                if product.main_image_url:
                    try:
                        from django.core.files.storage import default_storage
                        old_file_path = f'products/{product.main_image_url}'
                        if default_storage.exists(old_file_path):
                            default_storage.delete(old_file_path)
                    except:
                        pass
                
                # Создаем безопасное имя файла
                file_extension = os.path.splitext(product_image.name)[1]
                filename = f"{slugify(title)}_{user_id}_{timezone.now().timestamp()}{file_extension}"
                
                # Сохраняем файл
                from django.core.files.storage import default_storage
                from django.core.files.base import ContentFile
                
                file_path = default_storage.save(f'products/{filename}', ContentFile(product_image.read()))
                product.main_image_url = filename
            # ===== КОНЕЦ НОВОГО КОДА =====
            
            # Получаем новые товары из скрытого поля
            tovars_data = request.POST.get('tovars_data', '[]')
            new_tovar_texts = json.loads(tovars_data)
            
            # Валидация
            errors = []
            if not title or len(title.strip()) < 5:
                errors.append('Название товара должно содержать минимум 5 символов')
            if not description or len(description.strip()) < 10:
                errors.append('Описание должно содержать минимум 10 символов')
            if not price or float(price) <= 0:
                errors.append('Цена должна быть положительным числом')
            
            if errors:
                for error in errors:
                    messages.error(request, error)
                return render(request, 'user/edit_product.html', {
                    'user_obj': user,
                    'product': product,
                    'categories': Categories.objects.all(),
                    'product_types': ProductTypes.objects.all(),
                    'form_data': request.POST
                })
            
            # Обновляем продукт
            old_title = product.title
            product.title = title.strip()
            product.description = description.strip()
            product.price = price
            product.category_id = category_id if category_id else None
            product.product_type_id = product_type_id
            
            product.save()
            
            # Получаем существующие товары (только те, что не проданы)
            existing_tovars = list(product.tovars.filter(is_sold=False))
            existing_texts = [t.tovar_text for t in existing_tovars]
            
            # Добавляем только новые товары
            for text in new_tovar_texts:
                if text.strip() and text.strip() not in existing_texts:
                    tovar = Tovars.objects.create(
                        tovar_text=text.strip(),
                        is_sold=False
                    )
                    ProductsTovars.objects.create(
                        product=product,
                        tovar=tovar
                    )
            
            # Логирование
            UserActivityLog.objects.create(
                user=user,
                action='edit_product',
                description=f'Отредактирован товар: {old_title} -> {product.title}'
            )
            
            messages.success(request, f'Товар "{product.title}" успешно обновлен!')
            return redirect('user_products')
            
        except Exception as e:
            logger.error(f'Edit product error: {str(e)}')
            messages.error(request, f'Ошибка при редактировании товара: {str(e)}')
            return render(request, 'user/edit_product.html', {
                'user_obj': user,
                'product': product,
                'categories': Categories.objects.all(),
                'product_types': ProductTypes.objects.all(),
                'form_data': request.POST
            })
    
    return render(request, 'user/edit_product.html', {
        'user_obj': user,
        'product': product,
        'categories': Categories.objects.all(),
        'product_types': ProductTypes.objects.all()
    })

@csrf_exempt
@require_POST
def delete_product(request, product_id):
    """Удаление товара (API)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        # Проверяем, что товар принадлежит пользователю
        product = get_object_or_404(Products, id_product=product_id, seller=user)
        
        # Проверяем, нет ли активных заказов для этого товара
        has_active_orders = OrderItems.objects.filter(
            product=product,
            order__status__in=['pending', 'paid']
        ).exists()
        
        if has_active_orders:
            return JsonResponse({
                'error': 'Невозможно удалить товар, так как есть активные заказы'
            }, status=400)
        
        product_title = product.title
        
        # Помечаем товар как неактивный
        product.is_active = False
        product.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='delete_product',
            description=f'Удален товар: {product_title}'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Товар "{product_title}" успешно удален'
        })
        
    except Products.DoesNotExist:
        return JsonResponse({'error': 'Товар не найден'}, status=404)
    except Exception as e:
        logger.error(f'Delete product error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def delete_tovar(request):
    """Удаление товара (Tovar) из продукта"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        data = json.loads(request.body)
        tovar_id = data.get('tovar_id')
        product_id = data.get('product_id')
        
        # Проверяем, что товар и продукт существуют
        tovar = get_object_or_404(Tovars, id_tovar=tovar_id)
        product = get_object_or_404(Products, id_product=product_id, seller=user)
        
        # Проверяем, что товар не продан
        if tovar.is_sold:
            return JsonResponse({
                'error': 'Невозможно удалить проданный товар'
            }, status=400)
        
        # Удаляем связь из ProductsTovars
        ProductsTovars.objects.filter(product=product, tovar=tovar).delete()
        
        # Если товар больше нигде не используется, удаляем его
        if not tovar.products.exists():
            tovar.delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='delete_product',
            description=f'Удален товар #{tovar_id} из продукта {product.title}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Товар успешно удален'
        })
        
    except Exception as e:
        logger.error(f'Delete tovar error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

def product_stats(request, product_id):
    """Статистика по товару"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для просмотра статистики необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Проверяем, что товар принадлежит пользователю
    product = get_object_or_404(Products, id_product=product_id, seller=user)
    
    # Статистика
    stats = {
        'total_tovars': product.tovars.count(),
        'available_tovars': product.available_tovars_count,
        'sold_tovars': product.tovars.filter(is_sold=True).count(),
        'total_revenue': OrderItems.objects.filter(
            product=product,
            order__status__in=['paid', 'completed']
        ).aggregate(total=Sum('price_at_time_of_purchase'))['total'] or 0,
        'total_orders': OrderItems.objects.filter(product=product).count(),
        'avg_rating': ProductReviews.objects.filter(product=product).aggregate(avg=Avg('rating'))['avg'] or 0,
        'reviews_count': ProductReviews.objects.filter(product=product).count(),
    }
    
    # Последние продажи
    recent_sales = OrderItems.objects.filter(
        product=product
    ).select_related('order__user').order_by('-order__order_created_at')[:10]
    
    # График продаж за последние 30 дней
    end_date = timezone.now().date()
    start_date = end_date - timedelta(days=30)
    
    daily_sales = []
    for i in range(30):
        date = end_date - timedelta(days=i)
        daily_count = OrderItems.objects.filter(
            product=product,
            order__order_created_at__date=date,
            order__status__in=['paid', 'completed']
        ).count()
        
        daily_revenue = OrderItems.objects.filter(
            product=product,
            order__order_created_at__date=date,
            order__status__in=['paid', 'completed']
        ).aggregate(total=Sum('price_at_time_of_purchase'))['total'] or 0
        
        daily_sales.append({
            'date': date.strftime('%d.%m'),
            'count': daily_count,
            'revenue': float(daily_revenue)
        })
    
    daily_sales.reverse()
    
    return render(request, 'user/product_stats.html', {
        'user_obj': user,
        'product': product,
        'stats': stats,
        'recent_sales': recent_sales,
        'daily_sales': json.dumps(daily_sales)
    })




# ==================== ЧАТЫ ====================

def chat_buyer(request, chat_id=None):
    """Чат для покупателя"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Получаем все чаты пользователя как покупателя
    buyer_chats = Chats.objects.filter(
        buyer=user,
        is_active=True
    ).select_related('product', 'seller').order_by('-last_message_at')
    
    current_chat = None
    messages_list = []
    
    if chat_id:
        # Получаем конкретный чат
        current_chat = get_object_or_404(Chats, id_chat=chat_id, buyer=user)
        
        # Помечаем сообщения как прочитанные
        Messages.objects.filter(
            chat=current_chat,
            sender=current_chat.seller,
            is_read=False
        ).update(is_read=True, read_at=timezone.now())
        
        
        current_chat.save()
        
        # Получаем сообщения
        messages_list = Messages.objects.filter(
            chat=current_chat
        ).select_related('sender').order_by('sent_at')
    
    # Создание нового чата (если передан product_id)
    product_id = request.GET.get('product_id')
    if product_id and not chat_id:
        product = get_object_or_404(Products, id_product=product_id)
        
        # Проверяем, существует ли уже чат
        existing_chat = Chats.objects.filter(
            product=product,
            buyer=user,
            seller=product.seller
        ).first()
        
        if existing_chat:
            return redirect(f'/chat/buyer/{existing_chat.id_chat}/')
        
        # Создаем новый чат
        new_chat = Chats.objects.create(
            product=product,
            buyer=user,
            seller=product.seller,
            is_active=True
        )
        
        # Первое сообщение (приветственное)
        Messages.objects.create(
            chat=new_chat,
            sender=user,
            message_text=f'Здравствуйте! У меня вопрос по товару "{product.title}"'
        )
        
        return redirect(f'/chat/buyer/{new_chat.id_chat}/')
    
    return render(request, 'chat/chat_buyer.html', {
        'user_obj': user,
        'buyer_chats': buyer_chats,
        'current_chat': current_chat,
        'messages': messages_list
        
    })

def chat_seller(request, chat_id=None):
    """Чат для продавца"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Проверяем, есть ли у пользователя товары
    if not Products.objects.filter(seller=user).exists():
        messages.error(request, 'У вас нет товаров для продажи')
        return redirect('profile')
    
    # Получаем все чаты пользователя как продавца
    seller_chats = Chats.objects.filter(
        seller=user,
        is_active=True
    ).select_related('product', 'buyer').order_by('-last_message_at')
    
    current_chat = None
    messages_list = []
    
    if chat_id:
        # Получаем конкретный чат
        current_chat = get_object_or_404(Chats, id_chat=chat_id, seller=user)
        
        # Помечаем сообщения как прочитанные
        Messages.objects.filter(
            chat=current_chat,
            sender=current_chat.buyer,
            is_read=False
        ).update(is_read=True, read_at=timezone.now())
        
        
        current_chat.save()
        
        # Получаем сообщения
        messages_list = Messages.objects.filter(
            chat=current_chat
        ).select_related('sender').order_by('sent_at')
    
    return render(request, 'chat/chat_seller.html', {
        'user_obj': user,
        'seller_chats': seller_chats,
        'current_chat': current_chat,
        'messages': messages_list
        
    })


# Добавьте эти функции в конец файла views.py

@csrf_exempt
@require_POST
def create_dispute_chat(request, order_id):
    """Создание чата для оспаривания заказа (только один раз)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    order = get_object_or_404(Orders, id_order=order_id, user_id=user_id)
    
    # Проверяем, что заказ завершён (completed)
    if order.status != 'completed':
        return JsonResponse({'error': 'Оспорить можно только завершённый заказ'}, status=400)
    
    # Проверяем, нет ли уже чата по этому заказу (даже закрытого)
    existing_chat = Chats.objects.filter(order=order).first()
    if existing_chat:
        return JsonResponse({
            'success': True, 
            'chat_id': existing_chat.id_chat,
            'is_active': existing_chat.is_active,
            'message': 'Чат уже существует'
        })
    
    # Находим продавца и товары в заказе
    order_items = OrderItems.objects.filter(order=order).select_related('product__seller')
    seller = None
    product_titles = []
    
    if order_items.exists():
        seller = order_items.first().product.seller
        # Собираем названия всех товаров в заказе
        for item in order_items:
            if item.product.title not in product_titles:
                product_titles.append(item.product.title)
    
    # Формируем список товаров для сообщения
    products_list = ', '.join(product_titles)
    if len(product_titles) > 3:
        products_list = ', '.join(product_titles[:3]) + f' и ещё {len(product_titles) - 3} товаров'
    
    # Создаём чат
    chat = Chats.objects.create(
        product=None,
        buyer=user,
        seller=seller,
        order=order,
        is_active=True
    )
    
    # Отправляем системное сообщение с указанием товаров
    message_text = f'Создан спор по заказу №{order.id_order}\n📦 Товар(ы): {products_list}'
    Messages.objects.create(
        chat=chat,
        sender=user,
        message_text=message_text
    )
    
    # Логирование
    UserActivityLog.objects.create(
        user=user,
        action='send_message',
        description=f'Создан чат оспаривания для заказа #{order.id_order} (товары: {products_list})'
    )
    
    return JsonResponse({
        'success': True,
        'chat_id': chat.id_chat,
        'is_active': True,
        'message': 'Чат для оспаривания создан'
    })


@csrf_exempt
@require_POST
def close_dispute_chat(request, chat_id):
    """Закрытие чата оспаривания (только для менеджера)"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    # Только менеджер может закрыть спор
    if user.role_id != 3:
        return JsonResponse({'error': 'Только менеджер может завершить спор'}, status=403)
    
    chat = get_object_or_404(Chats, id_chat=chat_id)
    
    # Закрываем чат
    chat.is_active = False
    chat.save()
    
    # Отправляем системное сообщение
    Messages.objects.create(
        chat=chat,
        sender=user,
        message_text=f'Спор по заказу завершён менеджером {user.login}. Чат закрыт.'
    )
    
    # Логирование
    UserActivityLog.objects.create(
        user=user,
        action='admin_action',
        description=f'Закрыт чат оспаривания #{chat_id} по заказу #{chat.order.id_order if chat.order else "?"}'
    )
    
    return JsonResponse({'success': True, 'message': 'Чат оспаривания закрыт'})


@csrf_exempt
@require_POST
def send_message(request):
    """Отправка сообщения"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    try:
        data = json.loads(request.body)
        chat_id = data.get('chat_id')
        message_text = data.get('message_text')
        
        if not message_text or not message_text.strip():
            return JsonResponse({'error': 'Сообщение не может быть пустым'}, status=400)
        
        chat = get_object_or_404(Chats, id_chat=chat_id)
        sender = user
        
        # Проверяем, что пользователь является участником чата ИЛИ менеджером
        is_participant = (chat.buyer_id == user_id or chat.seller_id == user_id)
        is_manager = (user.role_id == 3)  # Менеджер
        
        if not is_participant and not is_manager:
            return JsonResponse({'error': 'Нет доступа к чату'}, status=403)
        
        # Если менеджер отправляет сообщение, он становится seller'ом чата (если seller не задан)
        if is_manager and chat.seller_id is None:
            chat.seller_id = user_id
            chat.save()
        
        # Создаем сообщение
        message = Messages.objects.create(
            chat=chat,
            sender=sender,
            message_text=message_text.strip()
        )
        
        # Обновляем время последнего сообщения
        chat.last_message_at = timezone.now()
        chat.save()
        
        # Если чат был закрыт - открываем его снова
        if not chat.is_active:
            chat.is_active = True
            chat.save()
        
        # Отправляем уведомления менеджерам (если отправитель - покупатель)
        if chat.buyer_id == user_id:
            notify_vk_managers(chat, message, sender)
        
        # Логирование
        UserActivityLog.objects.create(
            user_id=user_id,
            action='send_message',
            description=f'Отправлено сообщение в чате #{chat_id}'
        )
        
        return JsonResponse({
            'success': True,
            'message_id': message.id_message,
            'sender_name': message.sender.login,
            'sent_at': message.sent_at.strftime('%H:%M'),
            'message_text': message.message_text
        })
        
    except Exception as e:
        logger.error(f'Send message error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_GET
def get_messages(request, chat_id):
    """Получение сообщений чата"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        chat = get_object_or_404(Chats, id_chat=chat_id)
        
        # Проверяем доступ
        if chat.buyer_id != user_id and chat.seller_id != user_id:
            return JsonResponse({'error': 'Нет доступа к чату'}, status=403)
        
        messages = Messages.objects.filter(
            chat=chat
        ).select_related('sender').order_by('sent_at')
        
        serializer = MessagesSerializer(messages, many=True)
        
        
        
        chat.save()
        
        return JsonResponse({
            'success': True,
            'messages': serializer.data,
            'chat_info': {
                'product_title': chat.product.title if chat.product else 'Без товара',
                'buyer_name': chat.buyer.login,
                'seller_name': chat.seller.login
            }
        })
        
    except Exception as e:
        logger.error(f'Get messages error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_POST
def close_chat(request):
    """Закрытие чата"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    
    try:
        data = json.loads(request.body)
        chat_id = data.get('chat_id')
        
        chat = get_object_or_404(Chats, id_chat=chat_id)
        
        # Проверяем, что пользователь является участником чата
        if chat.buyer_id != user_id and chat.seller_id != user_id:
            return JsonResponse({'error': 'Нет доступа к чату'}, status=403)
        
        chat.is_active = False
        chat.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user_id=user_id,
            action='admin_action',
            description=f'Закрытие чата #{chat_id}'
        )
        
        return JsonResponse({'success': True, 'message': 'Чат закрыт'})
        
    except Exception as e:
        logger.error(f'Close chat error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

# ==================== ПОИСК ====================

def search_view(request):
    """Поиск товаров"""
    query = request.GET.get('q', '').strip()
    
    if not query:
        return redirect('products_list')
    
    products = Products.objects.filter(
        Q(title__icontains=query) | 
        Q(description__icontains=query),
        is_active=True
    ).select_related('seller', 'category', 'product_type')
    
    # Пагинация
    paginator = Paginator(products, 12)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Логирование поиска
    if 'user_id' in request.session:
        UserActivityLog.objects.create(
            user_id=request.session['user_id'],
            action='search',
            description=f'Поисковый запрос: "{query}"',
            additional_data={'query': query, 'results_count': products.count()}
        )
    
    return render(request, 'search/results.html', {
        'products': page_obj,
        'query': query,
        'results_count': products.count(),
        'page_obj': page_obj
    })

# ==================== УТИЛИТЫ ====================

def get_client_ip(request):
    """Получение IP адреса клиента"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip

# ==================== API VIEWSETS ====================

class ProductViewSet(viewsets.ModelViewSet):
    """API для товаров"""
    queryset = Products.objects.filter(is_active=True).select_related('seller', 'category', 'product_type')
    serializer_class = ProductsSerializer
    
    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Фильтрация
        category_id = self.request.query_params.get('category_id')
        if category_id:
            queryset = queryset.filter(category_id=category_id)
        
        product_type_id = self.request.query_params.get('product_type_id')
        if product_type_id:
            queryset = queryset.filter(product_type_id=product_type_id)
        
        min_price = self.request.query_params.get('min_price')
        if min_price:
            queryset = queryset.filter(price__gte=min_price)
        
        max_price = self.request.query_params.get('max_price')
        if max_price:
            queryset = queryset.filter(price__lte=max_price)
        
        # Поиск
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) | 
                Q(description__icontains=search)
            )
        
        # Сортировка
        sort_by = self.request.query_params.get('sort_by', '-rating')
        if sort_by in ['price', '-price', 'rating', '-rating', 'created_at', '-created_at']:
            queryset = queryset.order_by(sort_by)
        
        return queryset
    
    
    
    

class CartViewSet(viewsets.ModelViewSet):
    """API для корзины"""
    serializer_class = CartSerializer
    
    def get_queryset(self):
        if 'user_id' not in self.request.session:
            return Cart.objects.none()
        return Cart.objects.filter(user_id=self.request.session['user_id']).select_related('product')
    
    def create(self, request):
        if 'user_id' not in request.session:
            return Response({'error': 'Не авторизован'}, status=status.HTTP_401_UNAUTHORIZED)
        
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            serializer.save(user_id=request.session['user_id'])
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Сводка по корзине"""
        if 'user_id' not in request.session:
            return Response({'error': 'Не авторизован'}, status=status.HTTP_401_UNAUTHORIZED)
        
        cart_items = self.get_queryset()
        total_price = sum(item.product.price * item.quantity for item in cart_items)
        total_items = sum(item.quantity for item in cart_items)
        
        return Response({
            'total_price': float(total_price),
            'total_items': total_items,
            'items_count': cart_items.count()
        })
    
    @action(detail=False, methods=['delete'])
    def clear(self, request):
        """Очистка корзины"""
        if 'user_id' not in request.session:
            return Response({'error': 'Не авторизован'}, status=status.HTTP_401_UNAUTHORIZED)
        
        Cart.objects.filter(user_id=request.session['user_id']).delete()
        
        # Логирование
        UserActivityLog.objects.create(
            user_id=request.session['user_id'],
            action='remove_from_cart',
            description='Очистка корзины'
        )
        
        return Response({'success': True, 'message': 'Корзина очищена'})

class OrderViewSet(viewsets.ModelViewSet):
    """API для заказов"""
    serializer_class = OrdersSerializer
    
    def get_queryset(self):
        if 'user_id' not in self.request.session:
            return Orders.objects.none()
        
        user_id = self.request.session['user_id']
        user = Users.objects.get(id_user=user_id)
        
        # Админы и менеджеры видят все заказы
        if user.role_id in [1, 3]:  # Админ или менеджер
            return Orders.objects.all().select_related('user').order_by('-order_created_at')
        
        # Обычные пользователи видят только свои
        return Orders.objects.filter(user_id=user_id).select_related('user').order_by('-order_created_at')
    
    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """Отмена заказа"""
        if 'user_id' not in request.session:
            return Response({'error': 'Не авторизован'}, status=status.HTTP_401_UNAUTHORIZED)
        
        order = self.get_object()
        user_id = request.session['user_id']
        user = Users.objects.get(id_user=user_id)
        
        # Проверяем права
        if order.user_id != user_id and user.role_id not in [1, 3]:
            return Response(
                {'error': 'Недостаточно прав'}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        if order.status not in ['pending', 'paid']:
            return Response(
                {'error': 'Невозможно отменить заказ в текущем статусе'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Возвращаем товары на склад
        order_items = OrderItems.objects.filter(order=order).select_related('product', 'product_item')
        for item in order_items:
            if item.product_item:
                item.product_item.is_sold = False
                item.product_item.sold_at = None
                item.product_item.order_item = None
                item.product_item.save()
            
            item.product.stock_quantity += item.quantity
            item.product.save()
        
        # Возвращаем средства
        if order.status == 'paid':
            user.balance += order.total_cost
            user.save()
            
            Transactions.objects.create(
                user=user,
                order=order,
                amount=order.total_cost,
                transaction_type='refund',
                status='completed',
                reference=f'CANCEL_ORDER_{order.id_order}'
            )
        
        # Меняем статус
        order.status = 'cancelled'
        order.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=user,
            action='create_order',
            description=f'Отмена заказа №{order.id_order}'
        )
        
        return Response({'success': True, 'message': 'Заказ отменен'})

# ==================== АДМИН ПАНЕЛЬ ====================

# appip/views.py - обновите функцию admin_dashboard

def admin_dashboard(request):
    """Админ панель"""
    if 'user_id' not in request.session:
        return redirect('login')
    
    user_id = request.session['user_id']
    user = Users.objects.get(id_user=user_id)
    
    if user.role_id not in [1, 3]:  # Только админ и менеджер
        messages.error(request, 'Недостаточно прав')
        return redirect('home')
    
    # Статистика
    today = timezone.now().date()
    
    stats = {
        'users_count': Users.objects.count(),
        'products_count': Products.objects.count(),
        'orders_today': Orders.objects.filter(order_created_at__date=today).count(),
        'revenue_today': Orders.objects.filter(
            order_created_at__date=today, 
            status__in=['paid', 'completed']
        ).aggregate(total=Sum('total_cost'))['total'] or 0,
        'revenue_month': Orders.objects.filter(
            order_created_at__month=today.month,
            order_created_at__year=today.year,
            status__in=['paid', 'completed']
        ).aggregate(total=Sum('total_cost'))['total'] or 0,
        'pending_orders': Orders.objects.filter(status='pending').count(),
    }
    
    # Последние заказы
    recent_orders = Orders.objects.select_related('user').order_by('-order_created_at')[:10]
    
    # Последние пользователи
    recent_users = Users.objects.order_by('-registration_date')[:10]
    
    # Для менеджера: получаем чаты
    manager_chats = []
    if user.role_id == 3:  # Менеджер
        from .models import ManagerChats
        manager_chats = ManagerChats.objects.filter(
            manager=user,
            is_active=True
        ).select_related('user', 'chat').order_by('-last_message_at')
    
    # График продаж за последние 30 дней
    end_date = today
    start_date = end_date - timedelta(days=30)
    
    daily_stats = []
    for i in range(30):
        date = end_date - timedelta(days=i)
        daily_revenue = Orders.objects.filter(
            order_created_at__date=date,
            status__in=['paid', 'completed']
        ).aggregate(total=Sum('total_cost'))['total'] or 0
        
        daily_orders = Orders.objects.filter(order_created_at__date=date).count()
        
        daily_stats.append({
            'date': date.strftime('%d.%m'),
            'revenue': float(daily_revenue),
            'orders': daily_orders
        })
    
    daily_stats.reverse()
    
    context = {
        'stats': stats,
        'recent_orders': recent_orders,
        'recent_users': recent_users,
        'daily_stats': json.dumps(daily_stats),
        'is_admin': user.role_id == 1,
        'is_manager': user.role_id == 3,
        'manager_chats': manager_chats if user.role_id == 3 else [],
        'current_user': user,
    }
    
    return render(request, 'admin/dashboard.html', context)




@csrf_exempt
@require_GET
def get_all_users_for_manager(request):
    """Получение всех пользователей и чатов для менеджера"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    if user.role_id != 3:
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    # Получаем всех пользователей (обычных)
    users = Users.objects.filter(
        role_id=2
    ).exclude(id_user=user_id).order_by('login')
    
    users_data = []
    for u in users:
        users_data.append({
            'id': u.id_user,
            'login': u.login,
            'name': f"{u.firstname} {u.surname}",
            'role': u.role.role_name,
            'registration_date': u.registration_date.strftime('%d.%m.%Y'),
            'is_active': u.is_active
        })
    
    # Получаем все активные чаты (включая чаты по спорам)
    chats = Chats.objects.filter(is_active=True).select_related('buyer', 'seller', 'order')
    
    chats_data = []
    for chat in chats:
        # Показываем все чаты, где есть order (споры) ИЛИ чаты где seller - текущий менеджер
        if chat.order or chat.seller_id == user_id:
            unread_count = Messages.objects.filter(
                chat=chat,
                is_read=False
            ).exclude(sender=user).count()
            
            # Получаем последнее сообщение
            last_message = Messages.objects.filter(chat=chat).order_by('-sent_at').first()
            last_message_time = last_message.sent_at.strftime('%d.%m.%Y %H:%M') if last_message else chat.created_at.strftime('%d.%m.%Y %H:%M')
            
            chats_data.append({
                'id': chat.id_chat,
                'user_id': chat.buyer.id_user,
                'user_login': chat.buyer.login,
                'user_name': f"{chat.buyer.firstname} {chat.buyer.surname}",
                'seller_id': chat.seller.id_user if chat.seller else None,
                'seller_login': chat.seller.login if chat.seller else None,
                'order_id': chat.order.id_order if chat.order else None,
                'last_message_at': last_message_time,
                'unread_count': unread_count
            })
    
    # Сортируем чаты по времени последнего сообщения (новые сверху)
    chats_data.sort(key=lambda x: x['last_message_at'], reverse=True)
    
    return JsonResponse({
        'success': True, 
        'users': users_data,
        'chats': chats_data
    })

@csrf_exempt
@require_POST
def manager_start_chat(request):
    """Начало чата менеджера с пользователем"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    manager = get_object_or_404(Users, id_user=user_id)
    
    if manager.role_id != 3:  # Только менеджер
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        data = json.loads(request.body)
        user_id_to_chat = data.get('user_id')
        message_text = data.get('message_text', '')
        
        user_to_chat = get_object_or_404(Users, id_user=user_id_to_chat)
        
        # Проверяем, есть ли уже чат
        manager_chat = ManagerChats.objects.filter(
            manager=manager,
            user=user_to_chat
        ).first()
        
        # Проверяем, есть ли обычный чат
        existing_chat = Chats.objects.filter(
            buyer=user_to_chat,
            seller=manager
        ).first()
        
        chat = existing_chat
        
        if not existing_chat:
            # Создаем новый чат
            chat = Chats.objects.create(
                product=None,
                buyer=user_to_chat,
                seller=manager,
                is_active=True
            )
        
        if not manager_chat:
            # Создаем запись менеджерского чата
            manager_chat = ManagerChats.objects.create(
                manager=manager,
                user=user_to_chat,
                chat=chat,
                is_active=True
            )
        else:
            # Обновляем существующий
            manager_chat.chat = chat
            manager_chat.is_active = True
            manager_chat.save()
        
        # Если есть текст сообщения, отправляем его
        if message_text and message_text.strip():
            Messages.objects.create(
                chat=chat,
                sender=manager,
                message_text=message_text.strip()
            )
            
            # Обновляем время последнего сообщения
            chat.last_message_at = timezone.now()
            chat.save()
            
            manager_chat.last_message_at = timezone.now()
            manager_chat.save()
        
        # Логирование
        UserActivityLog.objects.create(
            user=manager,
            action='send_message',
            description=f'Менеджер начал чат с пользователем {user_to_chat.login}'
        )
        
        return JsonResponse({
            'success': True,
            'chat_id': chat.id_chat,
            'message': 'Чат успешно создан/найден'
        })
        
    except Exception as e:
        logger.error(f'Manager start chat error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_GET
def get_manager_chat_messages(request, chat_id):
    """Получение сообщений чата менеджера"""
    if 'user_id' not in request.session:
        return JsonResponse({'error': 'Не авторизован'}, status=401)
    
    user_id = request.session['user_id']
    user = get_object_or_404(Users, id_user=user_id)
    
    if user.role_id != 3:  # Только менеджер
        return JsonResponse({'error': 'Недостаточно прав'}, status=403)
    
    try:
        chat = get_object_or_404(Chats, id_chat=chat_id)
        
        # Если seller не задан, назначаем менеджера seller'ом чата
        if chat.seller_id is None:
            chat.seller_id = user_id
            chat.save()
        
        # Если чат был закрыт - показываем его менеджеру всё равно
        # (но при отправке сообщения он откроется)
        
        messages = Messages.objects.filter(
            chat=chat
        ).select_related('sender').order_by('sent_at')
        
        # Помечаем сообщения покупателя как прочитанные
        Messages.objects.filter(
            chat=chat,
            sender=chat.buyer,
            is_read=False
        ).update(is_read=True, read_at=timezone.now())
        
        serializer = MessagesSerializer(messages, many=True)
        
        return JsonResponse({
            'success': True,
            'messages': serializer.data,
            'chat_info': {
                'user_name': f"{chat.buyer.firstname} {chat.buyer.surname}",
                'user_login': chat.buyer.login,
                'user_id': chat.buyer.id_user,
                'order_id': chat.order.id_order if chat.order else None
            }
        })
        
    except Exception as e:
        logger.error(f'Get manager messages error: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

# ==================== ОБРАБОТЧИКИ ОШИБОК ====================

def handler404(request, exception):
    """Обработчик 404 ошибки"""
    return render(request, 'errors/404.html', status=404)

def handler500(request):
    """Обработчик 500 ошибки"""
    return render(request, 'errors/500.html', status=500)

# ==================== API ENDPOINTS ====================

@api_view(['GET'])
def api_categories(request):
    """API для получения категорий"""
    categories = Categories.objects.all()
    serializer = CategoriesSerializer(categories, many=True)
    return Response(serializer.data)

@api_view(['GET'])
def api_product_types(request):
    """API для получения типов товаров"""
    product_types = ProductTypes.objects.all()
    serializer = ProductTypesSerializer(product_types, many=True)
    return Response(serializer.data)

@api_view(['POST'])
def api_register(request):
    """API регистрации"""
    serializer = UsersSerializer(data=request.data)
    if serializer.is_valid():
        user = serializer.save()
        return Response({
            'success': True,
            'message': 'Регистрация успешна',
            'user_id': user.id_user
        }, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
def api_login(request):
    """API входа"""
    email = request.data.get('email')
    password = request.data.get('password')
    
    try:
        user = Users.objects.get(login=email, is_active=True)
        
        if user.check_password(password):
            # Здесь можно вернуть токен или установить сессию
            return Response({
                'success': True,
                'message': 'Вход успешен',
                'user': {
                    'id': user.id_user,
                    'email': user.login,
                    'name': f"{user.firstname} {user.surname}",
                    'role': user.role.role_name
                }
            })
        else:
            return Response(
                {'error': 'Неверный пароль'}, 
                status=status.HTTP_401_UNAUTHORIZED
            )
            
    except Users.DoesNotExist:
        return Response(
            {'error': 'Пользователь не найден'}, 
            status=status.HTTP_404_NOT_FOUND
        )

# Вспомогательные функции для работы с файлами
def download_receipt(request, order_id):
    """Скачивание чека заказа"""
    if 'user_id' not in request.session:
        messages.error(request, 'Для скачивания чека необходимо войти в систему')
        return redirect('login')
    
    user_id = request.session['user_id']
    order = get_object_or_404(Orders, id_order=order_id, user_id=user_id)
    
    # Создание документа Word
    document = Document()
    
    # Заголовок
    title = document.add_heading('Чек заказа', 0)
    title.alignment = 1
    
    # Информация о заказе
    document.add_heading('Информация о заказе', level=1)
    
    order_info = [
        ('Номер заказа', f'#{order.id_order}'),
        ('Дата заказа', order.order_created_at.strftime('%d.%m.%Y %H:%M')),
        ('Сумма заказа', f'{order.total_cost} ₽'),
        ('Статус', order.get_status_display()),
    ]
    
    for label, value in order_info:
        p = document.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)
    
    # Состав заказа
    document.add_heading('Состав заказа', level=1)
    
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Товар'
    hdr_cells[1].text = 'Количество'
    hdr_cells[2].text = 'Цена'
    hdr_cells[3].text = 'Сумма'
    
    order_items = OrderItems.objects.filter(order=order).select_related('product')
    for item in order_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.product.title
        row_cells[1].text = str(item.quantity)
        row_cells[2].text = f'{item.price_at_time_of_purchase} ₽'
        row_cells[3].text = f'{item.price_at_time_of_purchase * item.quantity} ₽'
    
    # Итог
    document.add_heading('Итоговая сумма', level=1)
    p = document.add_paragraph()
    p.add_run(f'Общая сумма: ').bold = True
    p.add_run(f'{order.total_cost} ₽')
    
    # Сохранение в буфер
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    # Логирование
    UserActivityLog.objects.create(
        user_id=user_id,
        action='download_receipt',
        description=f'Скачивание чека заказа №{order.id_order}'
    )
    
    # Отправка файла
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=receipt_order_{order_id}.docx'
    return response